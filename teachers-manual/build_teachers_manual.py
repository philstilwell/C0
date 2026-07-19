#!/usr/bin/env python3
"""Build, style, validate, and publish the Cø / N* teaching artifacts."""

from __future__ import annotations

from collections.abc import Callable
from dataclasses import dataclass
from datetime import datetime, timezone
import fcntl
import hashlib
import json
import os
from pathlib import Path
import platform
import re
import shutil
import subprocess
import sys
import tempfile
from zipfile import ZipFile

import docx
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
import pypdf
from pypdf import PdfReader


HERE = Path(__file__).resolve().parent
PROJECT = HERE.parent
MANUAL_SOURCE = HERE / "teachers-manual.md"
RESOURCE_SOURCE = HERE / "session-resource-pack.md"
BUILD_SCRIPT = Path(__file__).resolve()
TMP_ROOT = PROJECT / "tmp" / "docs" / "teachers-manual"
DOCX_DIR = PROJECT / "output" / "doc"
PDF_DIR = PROJECT / "output" / "pdf"
MANIFEST_OUT = PROJECT / "output" / "teachers-manual-build-manifest.json"

INSTRUCTOR_STEM = "teaching-c0-n-star-manual"
STUDENT_STEM = "c0-n-star-student-session-resource-pack"
REVEAL_STEM = "c0-n-star-instructor-controlled-reveal-sheets"
EDITION = "2.1"
CORPUS_BASELINE = "Core v2.0 (2026-07-08) plus repository companion manuscripts (2026-07-18)"
CORPUS_SOURCES = (
    PROJECT / "public" / "paper.pdf",
    PROJECT / "papers" / "where-is-the-conscious-subject" / "manuscript.md",
    PROJECT / "papers" / "consciousness-without-report" / "manuscript.md",
    PROJECT / "papers" / "indeterminacy-as-scientific-result" / "manuscript.md",
    PROJECT / "papers" / "ablating-n-star" / "manuscript.md",
    PROJECT / "papers" / "from-phenomenal-presence-to-phenomenal-character" / "manuscript.md",
    PROJECT / "papers" / "consciousness-in-the-schematic" / "manuscript.md",
)

BODY_FONT = os.environ.get("C0_BODY_FONT", "Liberation Serif")
DISPLAY_FONT = os.environ.get("C0_DISPLAY_FONT", "Liberation Sans")

NAVY = "1B344C"
BLUE = "2E6F8E"
PALE_BLUE = "EAF2F6"
PALE_GRAY = "F4F6F7"
WHITE = "FFFFFF"
TEXT = RGBColor(35, 42, 48)

TABLE_WIDTH_INCHES = 6.94

# Audited semantic column profiles. Values are relative weights and are normalized
# when written to the DOCX grid. Keeping every current signature explicit makes a
# new or renamed table fail closed until its content-to-width balance is reviewed.
TABLE_COLUMN_PROFILES: dict[tuple[str, ...], tuple[float, ...]] = {
    ("Part or session", "Focus", "PDF page"): (30, 62, 8),
    ("Element", "Role", "Question answered", "Characteristic false positive blocked"): (18, 20, 31, 31),
    ("Theory family", "What Cø / N* accepts", "What it rejects or tests", "Relation to the model"): (20, 23, 29, 28),
    ("Item", "Weight", "Purpose"): (30, 10, 60),
    ("Assessment", "Submission specification", "Schedule and permitted collaboration", "Feedback and revision rule"): (17, 28, 27, 28),
    ("Course outcome", "Principal instruction and practice", "Direct assessment evidence", "Minimum satisfactory evidence"): (9, 31, 23, 37),
    ("Meeting", "110-minute core sequence", "Material moved to prework or reference"): (21, 56, 23),
    ("Role", "Responsibility during the task", "Question the role must answer"): (22, 40, 38),
    ("Format", "Best use", "Procedure", "Required output"): (18, 20, 41, 21),
    ("Session", "Central question", "Main paper contribution"): (31, 35, 34),
    ("Time", "Teaching move"): (18, 82),
    ("Observation", "Directly supports", "Additional bridge needed for C0", "Rival interpretation to keep live"): (22, 22, 31, 25),
    ("Target", "Typical evidence", "Common mistake"): (20, 38, 42),
    ("N1", "N2", "N3", "N*", "Model prediction"): (9, 9, 9, 9, 64),
    ("Gate", "Estimate and interval", "Validity information", "Status"): (15, 28, 40, 17),
    ("Candidate", "Jself", "Jin", "AΘ", "Role complete?", "Initial result"): (22, 9, 9, 9, 17, 34),
    ("Question", "V", "N1"): (22, 39, 39),
    ("X", "Y", "Z"): (33, 33, 34),
    ("Estimator family", "What it may reveal", "Main limitation to teach"): (26, 31, 43),
    ("Tempting proxy", "Counterexample", "What N2 additionally requires"): (22, 39, 39),
    ("Recipient class", "Weight", "Effect and interval", "Latency", "Selectivity / route / context", "Class result"): (22, 8, 25, 10, 23, 12),
    ("Architecture", "How the state lasts", "Diagnostic perturbation", "What duration alone misses"): (21, 28, 27, 24),
    ("Cause", "What it means", "Proper next step"): (23, 34, 43),
    ("Feature", "Assessment A", "Assessment B"): (19, 33, 48),
    ("Type", "What is removed", "Inference"): (22, 31, 47),
    ("Claim", "Meaning", "Evidence required", "Scope of conclusion"): (17, 31, 31, 21),
    ("Presence result", "Bridge status", "Combined interpretation"): (21, 19, 60),
    ("Position", "Claim", "Explanatory status"): (23, 36, 41),
    ("Condition", "Competence question", "Possible partial failure"): (21, 40, 39),
    ("Case", "Initial component expectation", "Proper caution"): (24, 36, 40),
    ("Window", "Key observations", "Scientific output", "Cause profile and next step", "Separate action note"): (14, 28, 16, 23, 19),
    ("Transport and target evidence", "Within-theory four-way result", "Field-level qualifier"): (40, 22, 38),
    ("Candidate V criterion", "What it captures", "Likely false inclusion or exclusion"): (24, 31, 45),
    ("Case", "Cø / N*", "IIT 4.0-oriented intrinsic-structure analysis", "GNWT version requiring sustained global ignition or access", "RPT version treating local sensory recurrence as sufficient"): (22, 15, 24, 22, 17),
    ("Branch", "Required foundation", "Achievement", "What branch failure leaves intact"): (16, 31, 25, 28),
    ("Level", "Meaning"): (30, 70),
    ("Feature", "Weak design", "Stronger design"): (18, 30, 52),
    ("Phase", "Scientific aim", "Success criterion", "What success does not yet establish"): (21, 29, 31, 19),
    ("Argument", "Collaborative task", "Required artifact", "Validation or differentiation payoff", "Prohibited overreach"): (15, 16, 21, 27, 21),
    ("Objection", "Best available reply", "Residual pressure to acknowledge"): (20, 45, 35),
    ("Source", "Concept or commitment", "Primary treatment"): (23, 54, 23),
    ("Criterion", "4 - Complete", "3 - Proficient", "2 - Developing", "0-1 - Insufficient"): (12, 27, 19, 20, 22),
    ("Session", "Minimum satisfactory evidence", "Immediate response when missing"): (9, 55, 36),
    ("Dimension", "Strong", "Adequate", "Needs revision"): (15, 38, 20, 27),
    ("Track", "Required emphasis", "Conditionally applicable fields", "Minimum complete product"): (14, 38, 21, 27),
    ("Model", "Prospective prediction", "Expected anchor pattern", "Revision if disfavored"): (15, 30, 29, 26),
    ("Component", "Estimator family", "Threshold and uncertainty", "Main confound", "Negative control"): (12, 22, 27, 19, 20),
    ("Partition", "Permitted use", "Prohibited use"): (18, 46, 36),
    ("Project-appropriate outcome", "Exact rule", "Scientific interpretation", "Next study"): (23, 24, 31, 22),
    ("Checkpoint", "Team submission", "Reviewing team’s required response"): (18, 31, 51),
    ("Dimension", "Points", "Full-credit standard"): (24, 8, 68),
    ("Field", "Entry"): (38, 62),
    ("Question", "Cø / N*", "Rival theory"): (28, 36, 36),
    ("Step", "What to include", "Team record"): (20, 47, 33),
    ("Question", "Entry"): (40, 60),
    ("Condition", "Standard to inspect", "Decision and case-specific evidence"): (25, 42, 33),
    ("Evidence round", "Provisional output", "What changed?", "System change or license change?", "Next discriminating evidence"): (13, 18, 17, 27, 25),
    ("Criticism", "Source role or team", "Fatal, strengthening, or optional", "Accepted?", "Protocol change or reason for rejection"): (17, 15, 22, 12, 34),
    ("Term or symbol", "Teaching definition", "Guardrail"): (20, 41, 39),
    ("Corpus item", "Baseline date or version", "Stable project source"): (38, 30, 32),
    ("Team size", "Assignment rule"): (14, 86),
    ("Card", "Most direct target", "Possible relevance to C0", "Required bridge", "Strongest alternative", "Licensed conclusion", "Forbidden stronger inference"): (8, 13, 14, 13, 16, 15, 21),
    ("Card", "Most direct target", "Acceptable licensed conclusion", "Forbidden inference"): (7, 25, 35, 33),
    ("Profile", "Referent / bearer status", "V", "N1", "N2", "N3", "Additional information"): (8, 20, 13, 13, 13, 13, 20),
    ("Profile", "Candidate evidence", "Replacement and subset evidence"): (20, 36, 44),
    ("Synthetic intervention", "RC / mapping status", "V", "N1", "N2", "N3", "Independent anchor change"): (17, 11, 13, 13, 13, 13, 20),
    ("Intervention", "ΔV", "ΔN1", "ΔN2", "ΔN3"): (18, 20.5, 20.5, 20.5, 20.5),
    ("Field", "Team entry"): (42, 58),
    ("Class", "Weight", "Effect", "Latency", "Selective", "Route", "Context"): (15, 10, 17, 10, 12, 26, 10),
    ("Class", "Type", "Weight", "Effect", "Latency", "Selective", "Route", "Context"): (16, 14, 8, 20, 10, 12, 10, 10),
    ("Synthetic condition", "Interval", "Fidelity", "V", "N1", "N2", "N3", "Report accuracy"): (16, 12, 12, 12, 12, 12, 12, 12),
    ("Synthetic condition", "ΔV", "ΔN1", "ΔN2", "ΔN3", "RC / mapping status"): (18, 16, 16, 16, 16, 18),
    ("Row", "Weight", "V", "N1", "N2", "N3", "R*", "R-2", "Additional disclosed fact"): (5, 7, 10, 10, 10, 10, 7, 7, 34),
    ("", "R", "O", "G", "B"): (12, 22, 22, 22, 22),
    ("", "R", "O", "G"): (16, 28, 28, 28),
    ("Baseline phenomenal distances for B: to R, O, G", "Held-out rank accuracy", "Intervention prediction for R-O", "Observed R-O after intervention", "Validity"): (20, 12, 20, 20, 28),
    ("Trial-level legibility property", "Decision", "Evidence from the blind trial"): (36, 14, 50),
    ("Window", "Evidence"): (13, 87),
    ("Card", "Frozen source", "Required passage"): (7, 34, 59),
    ("Card", "Exact teaching version", "Primary target", "Bearer and central commitment", "Report and time", "Main caution"): (6, 18, 15, 24, 19, 18),
    ("Card", "Teaching version", "Primary pressure", "Comparison caution"): (6, 30, 34, 30),
    ("Compact capstone field", "Frozen team entry"): (42, 58),
    ("Enrollment", "Format"): (14, 86),
    ("Session", "Core resource in this appendix", "Optional existing packet", "Required worksheet"): (7, 36, 25, 32),
}

TABLE_CONTEXT_COLUMN_PROFILES: dict[
    tuple[tuple[str, ...], str], tuple[float, ...]
] = {
    (("Field", "Entry"), "Target form of consciousness"): (35, 65),
}

COMPACT_TABLE_WIDTHS: dict[tuple[str, ...], float] = {
    ("X", "Y", "Z"): 5.5,
    ("", "R", "O", "G", "B"): 5.5,
    ("", "R", "O", "G"): 5.5,
}

DENSE_TABLE_FONT_SIZES: dict[tuple[str, ...], float] = {
    ("Synthetic condition", "Interval", "Fidelity", "V", "N1", "N2", "N3", "Report accuracy"): 8.2,
    ("Synthetic condition", "ΔV", "ΔN1", "ΔN2", "ΔN3", "RC / mapping status"): 8.2,
    ("Row", "Weight", "V", "N1", "N2", "N3", "R*", "R-2", "Additional disclosed fact"): 9.0,
}

CENTERED_TABLE_COLUMNS: dict[tuple[str, ...], tuple[int, ...]] = {
    ("Part or session", "Focus", "PDF page"): (2,),
    ("Item", "Weight", "Purpose"): (1,),
    ("Course outcome", "Principal instruction and practice", "Direct assessment evidence", "Minimum satisfactory evidence"): (0,),
    ("N1", "N2", "N3", "N*", "Model prediction"): (0, 1, 2, 3),
    ("Gate", "Estimate and interval", "Validity information", "Status"): (3,),
    ("Candidate", "Jself", "Jin", "AΘ", "Role complete?", "Initial result"): (1, 2, 3, 4),
    ("X", "Y", "Z"): (0, 1, 2),
    ("Recipient class", "Weight", "Effect and interval", "Latency", "Selectivity / route / context", "Class result"): (1, 2, 3, 5),
    ("Session", "Minimum satisfactory evidence", "Immediate response when missing"): (0,),
    ("Dimension", "Points", "Full-credit standard"): (1,),
    ("Profile", "Referent / bearer status", "V", "N1", "N2", "N3", "Additional information"): (0, 2, 3, 4, 5),
    ("Synthetic intervention", "RC / mapping status", "V", "N1", "N2", "N3", "Independent anchor change"): (2, 3, 4, 5),
    ("Intervention", "ΔV", "ΔN1", "ΔN2", "ΔN3"): (0, 1, 2, 3, 4),
    ("Class", "Weight", "Effect", "Latency", "Selective", "Route", "Context"): (1, 2, 3, 4, 6),
    ("Class", "Type", "Weight", "Effect", "Latency", "Selective", "Route", "Context"): (2, 3, 4, 5, 6, 7),
    ("Synthetic condition", "Interval", "Fidelity", "V", "N1", "N2", "N3", "Report accuracy"): (1, 2, 3, 4, 5, 6, 7),
    ("Synthetic condition", "ΔV", "ΔN1", "ΔN2", "ΔN3", "RC / mapping status"): (1, 2, 3, 4, 5),
    ("Row", "Weight", "V", "N1", "N2", "N3", "R*", "R-2", "Additional disclosed fact"): (0, 1, 2, 3, 4, 5, 6, 7),
    ("", "R", "O", "G", "B"): (0, 1, 2, 3, 4),
    ("", "R", "O", "G"): (0, 1, 2, 3),
    ("Baseline phenomenal distances for B: to R, O, G", "Held-out rank accuracy", "Intervention prediction for R-O", "Observed R-O after intervention", "Validity"): (1, 2, 3),
    ("Trial-level legibility property", "Decision", "Evidence from the blind trial"): (1,),
    ("Card", "Frozen source", "Required passage"): (0,),
    ("Card", "Exact teaching version", "Primary target", "Bearer and central commitment", "Report and time", "Main caution"): (0,),
    ("Card", "Teaching version", "Primary pressure", "Comparison caution"): (0,),
    ("Enrollment", "Format"): (0,),
    ("Session", "Core resource in this appendix", "Optional existing packet", "Required worksheet"): (0,),
}

IMAGE_PATTERN = re.compile(r"!\[[^\]]*\]\(([^)]+)\)")
ANCHOR_PATTERN = re.compile(r"\{#([A-Za-z0-9_-]+)\}")
INTERNAL_LINK_PATTERN = re.compile(r"\]\(#([A-Za-z0-9_-]+)\)")
CONTENTS_ROW_PATTERN = re.compile(
    r"^\|\s*\[([^\]]+)\]\(#([A-Za-z0-9_-]+)\)\s*\|.*\|\s*(\d+)\s*\|$"
)
AUDIENCE_MARKER_PATTERN = re.compile(
    r"^<!-- audience:(shared|student|reveal|instructor) (begin|end) id=([a-z0-9-]+) -->$"
)
REVEAL_ID_PATTERN = re.compile(r"^session-(0[1-9]|1[0-4])-[a-z0-9-]+$")
EXPECTED_REVEAL_IDS = (
    "session-03-audit-1",
    "session-03-audit-2",
    "session-03-audit-3",
    "session-04-complication-1",
    "session-04-complication-2",
    "session-04-complication-3",
    "session-04-complication-4",
    "session-04-complication-5",
    "session-06-team-a",
    "session-06-team-b",
    "session-06-team-c",
    "session-06-team-d",
    "session-06-synthesis",
    "session-07-variant-a",
    "session-07-variant-b",
    "session-07-variant-c",
    "session-07-variant-d",
    "session-07-variant-e",
    "session-07-variant-f",
    "session-08-extension-rows",
    "session-08-mechanism-dossiers",
    "session-08-audit-1",
    "session-08-audit-2",
    "session-08-audit-3",
    "session-08-audit-4",
    "session-09-variant-a",
    "session-09-variant-b",
    "session-09-symmetry-extension",
    "session-10-profile-a",
    "session-10-profile-b",
    "session-10-profile-c",
    "session-10-profile-d",
    "session-10-profile-e",
    "session-11-stage-2",
    "session-11-stage-3",
    "session-11-longitudinal-extension",
    "session-12-dossier-a",
    "session-12-dossier-b",
    "session-12-dossier-c",
    "session-12-dossier-d",
    "session-13-profile-a",
    "session-13-profile-b",
    "session-13-profile-c",
    "session-13-profile-d",
    "session-13-profile-e",
    "session-14-result-1",
    "session-14-result-2",
    "session-14-result-3",
    "session-14-result-4",
    "session-14-result-5",
)
EXPECTED_STUDENT_BLOCK_IDS = tuple(
    block_id
    for session in range(1, 15)
    for block_id in (
        ((f"session-{session:02d}-base",) if session in {1, 2, 3, 5, 7} else (
            f"session-{session:02d}-base-a",
        ))
        + (() if session in {1, 2, 5} else (f"session-{session:02d}-base-b",))
        + (f"session-{session:02d}-access",)
    )
)
EXPECTED_INSTRUCTOR_BLOCK_IDS = (
    "appendix-h-intro",
    *(f"session-{session:02d}-key" for session in range(1, 15)),
    "instructor-preparation-and-closing",
)
EXPECTED_AUDIENCE_BY_ID = {
    "common-use-rules": "shared",
    **{block_id: "student" for block_id in EXPECTED_STUDENT_BLOCK_IDS},
    **{block_id: "reveal" for block_id in EXPECTED_REVEAL_IDS},
    **{block_id: "instructor" for block_id in EXPECTED_INSTRUCTOR_BLOCK_IDS},
}
_expected_block_order = ["appendix-h-intro", "common-use-rules"]
for _session in range(1, 15):
    if _session in {1, 2, 3, 5, 7}:
        _expected_block_order.append(f"session-{_session:02d}-base")
    else:
        _expected_block_order.append(f"session-{_session:02d}-base-a")
    _expected_block_order.extend(
        block_id
        for block_id in EXPECTED_REVEAL_IDS
        if block_id.startswith(f"session-{_session:02d}-")
    )
    if _session not in {1, 2, 5}:
        _expected_block_order.append(f"session-{_session:02d}-base-b")
    _expected_block_order.extend(
        (f"session-{_session:02d}-key", f"session-{_session:02d}-access")
    )
_expected_block_order.append("instructor-preparation-and-closing")
EXPECTED_BLOCK_ORDER = tuple(_expected_block_order)
del _expected_block_order, _session


@dataclass(frozen=True)
class AudienceBlock:
    audience: str
    block_id: str
    start_line: int
    end_line: int
    markdown: str


@dataclass(frozen=True)
class RevealSpec:
    block_id: str
    session: int
    title: str


@dataclass(frozen=True)
class ContentsReference:
    label: str
    anchor: str
    printed_page: int


def run(*args: str, cwd: Path | None = None) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        args,
        cwd=cwd,
        check=True,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )


def tool_output(*args: str) -> str:
    return run(*args).stdout.strip()


def find_tool(name: str, env_name: str) -> str:
    override = os.environ.get(env_name)
    if override:
        candidate = Path(override).expanduser()
        if not candidate.is_file() or not os.access(candidate, os.X_OK):
            raise FileNotFoundError(f"{env_name} does not name an executable file: {candidate}")
        return str(candidate)
    found = shutil.which(name)
    if found:
        return found
    raise FileNotFoundError(
        f"Required tool not found on PATH: {name}. Set {env_name} to its executable path."
    )


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def normalized_docx_content_sha256(path: Path) -> str:
    """Hash DOCX member content while excluding volatile core timestamps."""
    digest = hashlib.sha256()
    with ZipFile(path) as archive:
        for name in sorted(archive.namelist()):
            if name == "docProps/core.xml" or name.endswith("/"):
                continue
            digest.update(name.encode("utf-8"))
            digest.update(b"\0")
            digest.update(archive.read(name))
            digest.update(b"\0")
    return digest.hexdigest()


def normalized_pdf_content_sha256(path: Path) -> str:
    """Hash page geometry, drawing streams, text, links, and outline without PDF IDs/dates."""
    reader = PdfReader(str(path))
    digest = hashlib.sha256()
    page_by_id = {
        page.indirect_reference.idnum: index
        for index, page in enumerate(reader.pages)
        if page.indirect_reference is not None
    }
    for index, page in enumerate(reader.pages):
        digest.update(f"page:{index}:{tuple(float(v) for v in page.mediabox)}".encode("utf-8"))
        digest.update((page.extract_text() or "").encode("utf-8"))
        contents = page.get_contents()
        if contents is not None:
            digest.update(contents.get_data())
        for annotation in page.get("/Annots") or []:
            item = annotation.get_object()
            digest.update(str(item.get("/Subtype") or "").encode("utf-8"))
            digest.update(str(item.get("/Rect") or "").encode("utf-8"))
            destination = item.get("/Dest")
            if isinstance(destination, list) and destination:
                target_id = getattr(destination[0], "idnum", None)
                digest.update(f"dest:{page_by_id.get(target_id)}".encode("utf-8"))
    digest.update(json.dumps(flatten_outline_pages(reader), ensure_ascii=False).encode("utf-8"))
    return digest.hexdigest()


def normalized_content_sha256(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return normalized_docx_content_sha256(path)
    if path.suffix.lower() == ".pdf":
        return normalized_pdf_content_sha256(path)
    raise ValueError(f"No normalized-content digest is defined for {path}")


def font_file_report(soffice: str) -> dict[str, object]:
    """Record the exact bundled font programs used by the LibreOffice runtime."""
    executable = Path(soffice).resolve()
    search_root = executable.parents[2] if len(executable.parents) > 2 else executable.parent
    records: dict[str, object] = {}
    for role, family in (("body", BODY_FONT), ("display", DISPLAY_FONT)):
        family_stem = family.replace(" ", "")
        candidates = sorted(
            path
            for path in search_root.glob(
                f"native/libreoffice-headless/libreoffice/LibreOfficeDev.app/Contents/Resources/fonts/**/{family_stem}-*.ttf"
            )
            if path.is_file()
        )
        if not candidates:
            return {
                "checked": False,
                "reason": f"Bundled font files not found for {family!r} under {search_root}",
            }
        records[role] = {
            "family": family,
            "files": [
                {"path": str(path), "sha256": sha256(path)} for path in candidates
            ],
        }
    return {"checked": True, "files": records}


def source_inputs() -> list[Path]:
    """Return all build inputs and audited corpus dependencies for this release."""
    inputs = [
        MANUAL_SOURCE,
        RESOURCE_SOURCE,
        BUILD_SCRIPT,
        HERE / "README.md",
        HERE / "requirements-teachers-manual.txt",
        *CORPUS_SOURCES,
    ]
    for source in (MANUAL_SOURCE, RESOURCE_SOURCE):
        text = source.read_text(encoding="utf-8")
        for target in IMAGE_PATTERN.findall(text):
            clean_target = target.split(" ", 1)[0].strip("<>")
            if "://" in clean_target or clean_target.startswith("data:"):
                continue
            inputs.append((HERE / clean_target).resolve())
    return sorted(set(inputs))


def contents_references() -> list[ContentsReference]:
    """Parse the printed instructor contents table and its internal targets."""
    references: list[ContentsReference] = []
    for line in MANUAL_SOURCE.read_text(encoding="utf-8").splitlines():
        match = CONTENTS_ROW_PATTERN.fullmatch(line)
        if match:
            label, anchor, page = match.groups()
            references.append(ContentsReference(label, anchor, int(page)))
    if len(references) != 27:
        raise ValueError(f"Instructor contents inventory must contain 27 rows; found {len(references)}")
    if len({reference.anchor for reference in references}) != len(references):
        raise ValueError("Instructor contents inventory contains a duplicate anchor")
    return references


def validate_worksheet_introductions() -> int:
    """Require an explanatory paragraph immediately after each worksheet heading."""
    lines = MANUAL_SOURCE.read_text(encoding="utf-8").splitlines()
    verified = 0
    for number in range(1, 10):
        prefix = f"### Worksheet {number}."
        positions = [index for index, line in enumerate(lines) if line.startswith(prefix)]
        if len(positions) != 1:
            raise ValueError(f"Worksheet {number} heading count is {len(positions)}; expected one")
        for line in lines[positions[0] + 1 :]:
            if not line.strip():
                continue
            if line.startswith(("|", "#", "\\newpage")):
                raise ValueError(f"Worksheet {number} lacks an introductory explanation")
            verified += 1
            break
    return verified


def validate_sources(inputs: list[Path]) -> dict[str, int]:
    missing = [str(path) for path in inputs if not path.is_file()]
    if missing:
        raise FileNotFoundError(f"Missing build inputs: {missing}")
    combined = "\n".join(
        source.read_text(encoding="utf-8") for source in (MANUAL_SOURCE, RESOURCE_SOURCE)
    )
    anchors = ANCHOR_PATTERN.findall(combined)
    duplicates = sorted({anchor for anchor in anchors if anchors.count(anchor) > 1})
    if duplicates:
        raise ValueError(f"Duplicate explicit anchors: {duplicates}")
    unresolved = sorted(set(INTERNAL_LINK_PATTERN.findall(combined)) - set(anchors))
    if unresolved:
        raise ValueError(f"Internal links lack explicit targets: {unresolved}")
    audience_blocks = parse_audience_blocks()
    references = contents_references()
    return {
        "source_files": len(inputs),
        "corpus_source_files": len(CORPUS_SOURCES),
        "explicit_anchors": len(anchors),
        "internal_link_targets": len(INTERNAL_LINK_PATTERN.findall(combined)),
        "contents_references": len(references),
        "worksheet_introductions": validate_worksheet_introductions(),
        "local_figures": sum(path.suffix.lower() in {".png", ".jpg", ".jpeg", ".svg"} for path in inputs),
        "audience_blocks": len(audience_blocks),
        "controlled_reveal_blocks": sum(block.audience == "reveal" for block in audience_blocks),
    }


def set_font(style, name: str, size: float, *, bold=None, italic=None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), name)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def set_run_font(run, name: str, size: float | None = None) -> None:
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), name)


def set_math_run_format(paragraph, size: float, color: str) -> None:
    """Apply table typography to Office Math runs, which python-docx omits."""
    half_points = str(round(size * 2))
    for math_run in paragraph._p.xpath(".//m:r"):
        run_properties = math_run.find(qn("w:rPr"))
        if run_properties is None:
            run_properties = OxmlElement("w:rPr")
            math_properties = math_run.find(qn("m:rPr"))
            if math_properties is None:
                math_run.insert(0, run_properties)
            else:
                math_run.insert(math_run.index(math_properties) + 1, run_properties)
        for property_name in ("sz", "szCs"):
            property_node = run_properties.find(qn(f"w:{property_name}"))
            if property_node is None:
                property_node = OxmlElement(f"w:{property_name}")
                run_properties.append(property_node)
            property_node.set(qn("w:val"), half_points)
        color_node = run_properties.find(qn("w:color"))
        if color_node is None:
            color_node = OxmlElement("w:color")
            run_properties.append(color_node)
        color_node.set(qn("w:val"), color)


SUBSCRIPT_TRANSLATION = str.maketrans("0123456789+-=()", "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎")
SUPERSCRIPT_TRANSLATION = str.maketrans("0123456789+-=()", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾")
SCRIPT_NORMALIZATION = str.maketrans(
    "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾",
    "0123456789+-=()0123456789+-=()",
)


def office_math_text(element) -> str:
    """Flatten a small Office Math expression while retaining simple scripts."""
    if element.tag == qn("m:sSub"):
        base = element.find(qn("m:e"))
        subscript = element.find(qn("m:sub"))
        base_text = office_math_text(base) if base is not None else ""
        subscript_text = office_math_text(subscript) if subscript is not None else ""
        translated = subscript_text.translate(SUBSCRIPT_TRANSLATION)
        suffix = translated if translated != subscript_text or subscript_text.isdigit() else f"_{subscript_text}"
        return base_text + suffix
    if element.tag == qn("m:sSup"):
        base = element.find(qn("m:e"))
        superscript = element.find(qn("m:sup"))
        base_text = office_math_text(base) if base is not None else ""
        superscript_text = office_math_text(superscript) if superscript is not None else ""
        translated = superscript_text.translate(SUPERSCRIPT_TRANSLATION)
        suffix = translated if translated != superscript_text or superscript_text.isdigit() else f"^{superscript_text}"
        return base_text + suffix
    if element.tag == qn("m:t"):
        return element.text or ""
    return "".join(office_math_text(child) for child in element)


def flatten_office_math(container) -> None:
    """Replace math with styled text where compact table typography is essential."""
    for equation in list(container.xpath(".//m:oMath")):
        parent = equation.getparent()
        run = OxmlElement("w:r")
        text_node = OxmlElement("w:t")
        text_node.text = office_math_text(equation)
        run.append(text_node)
        parent.replace(equation, run)


def set_shading(element, fill: str) -> None:
    properties = element.get_or_add_tcPr() if hasattr(element, "get_or_add_tcPr") else element
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=70, bottom=70, start=85, end=85) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def normalized_table_text(cell) -> str:
    """Return visible cell text, including Office Math runs, for stable profiling."""
    pieces = [
        node.text
        for node in cell._tc.iter()
        if node.tag in {qn("w:t"), qn("m:t")} and node.text
    ]
    text = "".join(pieces).replace("\u00a0", " ").translate(SCRIPT_NORMALIZATION)
    text = text.replace("_", "").replace("^", "")
    for dash in ("\u2212", "\u2013", "\u2014"):
        text = text.replace(dash, "-")
    return re.sub(r"\s+", " ", text).strip()


def table_signature(table) -> tuple[str, ...]:
    if not table.rows:
        return ()
    return tuple(normalized_table_text(cell) for cell in table.rows[0].cells)


def table_layout_profile(table) -> tuple[tuple[str, ...], tuple[float, ...], float, float]:
    signature = table_signature(table)
    first_record = normalized_table_text(table.cell(1, 0)) if len(table.rows) > 1 else ""
    weights = TABLE_CONTEXT_COLUMN_PROFILES.get(
        (signature, first_record), TABLE_COLUMN_PROFILES.get(signature)
    )
    if weights is None:
        raise ValueError(f"Table lacks an audited column-width profile: {signature!r}")
    if len(weights) != len(table.columns):
        raise ValueError(
            f"Table profile has {len(weights)} columns but document has {len(table.columns)}: "
            f"{signature!r}"
        )
    if any(weight <= 0 for weight in weights):
        raise ValueError(f"Table profile contains a nonpositive width: {signature!r} {weights!r}")
    total_width = COMPACT_TABLE_WIDTHS.get(signature, TABLE_WIDTH_INCHES)
    font_size = DENSE_TABLE_FONT_SIZES.get(signature, 9.0)
    return signature, weights, total_width, font_size


def fixed_column_twips(weights: tuple[float, ...], total_width: float) -> list[int]:
    total_twips = round(total_width * 1440)
    weight_total = sum(weights)
    widths = [round(total_twips * weight / weight_total) for weight in weights]
    widths[-1] += total_twips - sum(widths)
    return widths


def set_fixed_table_columns(table, widths: list[int]) -> None:
    """Write preferred width, fixed layout, grid widths, and every cell width."""
    if len(widths) != len(table.columns):
        raise ValueError("Column-width count does not match the table")
    table.autofit = False
    table_properties = table._tbl.tblPr
    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")

    layout = table_properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid_columns = table._tbl.tblGrid.gridCol_lst
    if len(grid_columns) != len(widths):
        raise ValueError("DOCX table grid does not match the logical column count")
    for grid_column, width in zip(grid_columns, widths, strict=True):
        grid_column.set(qn("w:w"), str(width))

    for row in table.rows:
        if len(row.cells) != len(widths):
            raise ValueError("Merged table cells require a separately audited width profile")
        for cell, width in zip(row.cells, widths, strict=True):
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def validate_docx_table_layouts(document) -> dict[str, object]:
    full_width = 0
    compact = 0
    dense = 0
    for index, table in enumerate(document.tables):
        signature, weights, total_width, _ = table_layout_profile(table)
        expected = fixed_column_twips(weights, total_width)
        actual = [
            int(grid_column.get(qn("w:w")))
            for grid_column in table._tbl.tblGrid.gridCol_lst
        ]
        if actual != expected:
            raise ValueError(
                f"Table {index} grid differs from its audited profile: "
                f"{signature!r}; expected={expected}, actual={actual}"
            )
        layout = table._tbl.tblPr.find(qn("w:tblLayout"))
        if layout is None or layout.get(qn("w:type")) != "fixed":
            raise ValueError(f"Table {index} is not fixed-layout: {signature!r}")
        for row in table.rows:
            for column, cell in enumerate(row.cells):
                cell_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                if (
                    cell_width is None
                    or cell_width.get(qn("w:type")) != "dxa"
                    or int(cell_width.get(qn("w:w"))) != expected[column]
                ):
                    raise ValueError(
                        f"Table {index} cell width differs from its audited grid: {signature!r}"
                    )
        if signature in COMPACT_TABLE_WIDTHS:
            compact += 1
        else:
            full_width += 1
        if signature in DENSE_TABLE_FONT_SIZES:
            dense += 1
    return {
        "audited_tables": len(document.tables),
        "fixed_layout_tables": len(document.tables),
        "full_width_tables": full_width,
        "compact_tables": compact,
        "dense_table_profiles": dense,
        "default_width_inches": TABLE_WIDTH_INCHES,
    }


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def parse_audience_blocks() -> list[AudienceBlock]:
    """Parse the resource source with fail-closed audience classification."""
    lines = RESOURCE_SOURCE.read_text(encoding="utf-8").splitlines()
    blocks: list[AudienceBlock] = []
    seen_ids: set[str] = set()
    current: tuple[str, str, int] | None = None
    content: list[str] = []

    for line_number, line in enumerate(lines, start=1):
        marker = AUDIENCE_MARKER_PATTERN.fullmatch(line)
        if line.startswith("<!-- audience:") and marker is None:
            raise ValueError(f"Malformed audience marker at line {line_number}: {line}")
        if marker:
            audience, action, block_id = marker.groups()
            if action == "begin":
                if current is not None:
                    raise ValueError(f"Nested audience block at line {line_number}")
                if block_id in seen_ids:
                    raise ValueError(f"Duplicate audience block id: {block_id}")
                seen_ids.add(block_id)
                current = (audience, block_id, line_number)
                content = []
            else:
                if current is None:
                    raise ValueError(f"Orphan audience-block end at line {line_number}")
                open_audience, open_id, start_line = current
                if (audience, block_id) != (open_audience, open_id):
                    raise ValueError(
                        f"Mismatched audience-block end at line {line_number}: "
                        f"expected {open_audience}/{open_id}, got {audience}/{block_id}"
                    )
                markdown = "\n".join(content).strip()
                if not markdown:
                    raise ValueError(f"Empty audience block: {block_id}")
                blocks.append(
                    AudienceBlock(open_audience, open_id, start_line, line_number, markdown)
                )
                current = None
                content = []
            continue

        if current is not None:
            content.append(line)
        elif line.strip() not in {"", "---"}:
            raise ValueError(
                f"Substantive resource content is not audience-classified at line {line_number}: {line}"
            )

    if current is not None:
        raise ValueError(f"Unclosed audience block: {current[1]}")

    joined = "\n".join(block.markdown for block in blocks)
    base = "\n".join(
        block.markdown for block in blocks if block.audience in {"shared", "student"}
    )
    reveal = "\n".join(block.markdown for block in blocks if block.audience == "reveal")
    actual_audience_by_id = {block.block_id: block.audience for block in blocks}
    if actual_audience_by_id != EXPECTED_AUDIENCE_BY_ID:
        missing = sorted(set(EXPECTED_AUDIENCE_BY_ID) - set(actual_audience_by_id))
        unexpected = sorted(set(actual_audience_by_id) - set(EXPECTED_AUDIENCE_BY_ID))
        changed = sorted(
            block_id
            for block_id in set(actual_audience_by_id) & set(EXPECTED_AUDIENCE_BY_ID)
            if actual_audience_by_id[block_id] != EXPECTED_AUDIENCE_BY_ID[block_id]
        )
        raise ValueError(
            "Audience block inventory differs from the audited release: "
            f"missing={missing}, unexpected={unexpected}, audience_changed={changed}"
        )
    actual_order = tuple(block.block_id for block in blocks)
    if actual_order != EXPECTED_BLOCK_ORDER:
        raise ValueError("Audience block order differs from the audited release")
    key_blocks = [block for block in blocks if block.block_id.startswith("session-") and block.block_id.endswith("-key")]
    access_blocks = [block for block in blocks if block.block_id.startswith("session-") and block.block_id.endswith("-access")]
    if len(key_blocks) != 14 or any(block.audience != "instructor" for block in key_blocks) or joined.count("### Instructor Key") != 14:
        raise ValueError("Audience source must contain exactly 14 session answer-key blocks")
    if len(access_blocks) != 14 or any(block.audience != "student" for block in access_blocks):
        raise ValueError("Audience source must contain exactly 14 student accessibility blocks")
    if base.count("## Session ") != 14:
        raise ValueError("Student-base audience must contain exactly 14 session headings")
    if "Instructor Key" in base or "Reveal Sheet" in base:
        raise ValueError("Student-base audience contains a key or controlled reveal")
    if "Instructor Key" in reveal or "Instructor Preparation Crosswalk" in reveal:
        raise ValueError("Reveal audience contains instructor-only content")
    reveal_blocks = [block for block in blocks if block.audience == "reveal"]
    reveal_ids = tuple(block.block_id for block in reveal_blocks)
    if reveal_ids != EXPECTED_REVEAL_IDS:
        missing = sorted(set(EXPECTED_REVEAL_IDS) - set(reveal_ids))
        unexpected = sorted(set(reveal_ids) - set(EXPECTED_REVEAL_IDS))
        raise ValueError(
            "Controlled-reveal inventory or order differs from the audited release: "
            f"missing={missing}, unexpected={unexpected}"
        )
    reveal_specs = [reveal_spec(block) for block in reveal_blocks]
    if len({spec.title for spec in reveal_specs}) != len(reveal_specs):
        raise ValueError("Controlled-reveal titles must be unique")
    return blocks


def reveal_spec(block: AudienceBlock) -> RevealSpec:
    """Derive and validate release metadata from one controlled-reveal block."""
    if block.audience != "reveal":
        raise ValueError(f"Block is not a controlled reveal: {block.block_id}")
    id_match = REVEAL_ID_PATTERN.fullmatch(block.block_id)
    if id_match is None:
        raise ValueError(f"Reveal block has an invalid typed id: {block.block_id}")
    session = int(id_match.group(1))
    lines = block.markdown.splitlines()
    if not lines:
        raise ValueError(f"Reveal block is empty: {block.block_id}")
    heading_match = re.fullmatch(r"### (Reveal Sheet - Session ([1-9]|1[0-4]) .+)", lines[0])
    if heading_match is None:
        raise ValueError(
            f"Reveal block {block.block_id} must begin with one typed Session 1-14 heading"
        )
    heading_session = int(heading_match.group(2))
    if heading_session != session:
        raise ValueError(
            f"Reveal id/heading session mismatch for {block.block_id}: "
            f"id={session}, heading={heading_session}"
        )
    extra_headings = [line for line in lines[1:] if line.startswith("### Reveal Sheet - Session ")]
    if extra_headings:
        raise ValueError(f"Reveal block contains multiple release headings: {block.block_id}")
    return RevealSpec(block.block_id, session, heading_match.group(1))


def audience_markdown(kind: str, blocks: list[AudienceBlock]) -> tuple[str, list[str]]:
    """Render instructor, student-base, or controlled-reveal Markdown with provenance."""
    if kind == "instructor":
        selected = blocks
        header: list[str] = []
    elif kind == "student":
        selected = [block for block in blocks if block.audience in {"shared", "student"}]
        header = [
            "---",
            'title: "Cø / N* Student Base Resource Pack"',
            'subtitle: "Key-free and reveal-safe collaborative materials for Sessions 1-14"',
            'author: "Phil Stilwell"',
            'date: "Student base edition 2.1 - July 18, 2026"',
            "lang: en-US",
            "---",
            "",
            r"\newpage",
            "",
            "# Student Use Guide {#student-use}",
            "",
            "This pack contains initial student-facing materials only. It omits answer keys and every instructor-controlled reveal. Your instructor will release assigned evidence cards after the stated commitment point. All cases and values are synthetic and for teaching only; do not use them for clinical, welfare, or consciousness judgments outside the course.",
            "",
        ]
    elif kind == "reveal":
        selected = [block for block in blocks if block.audience == "reveal"]
        header = [
            "---",
            'title: "Cø / N* Instructor-Controlled Student Reveal Sheets"',
            'subtitle: "Key-free staged evidence cards - do not distribute as a complete student pack"',
            'author: "Phil Stilwell"',
            'date: "Controlled-reveal edition 2.1 - July 18, 2026"',
            "lang: en-US",
            "---",
            "",
            r"\newpage",
            "",
            "# Reveal-Sheet Use Guide {#reveal-use}",
            "",
            "These sheets contain no answer keys, but they include evidence that must remain hidden until a team reaches the stated commitment point. Instructors distribute only the assigned sheet or page, never this complete file. All cases and values are synthetic and for teaching only.",
            "",
        ]
    else:
        raise ValueError(f"Unknown audience artifact kind: {kind}")

    body: list[str] = []
    for index, block in enumerate(selected):
        markdown = block.markdown
        if kind == "reveal":
            markdown = re.sub(r"^### Reveal Sheet", "## Reveal Sheet", markdown, count=1)
            if index:
                body.extend(["", r"\newpage", ""])
        body.append(markdown)
        body.append("")
    result = "\n".join(header + body).rstrip() + "\n"
    if "<!-- audience:" in result:
        raise ValueError(f"Raw audience marker leaked into {kind} Markdown")
    return result, [block.block_id for block in selected]


def pandoc_docx(pandoc: str, sources: list[Path], destination: Path) -> None:
    command = [
        pandoc,
        *[str(source) for source in sources],
        "--standalone",
        "--from=markdown+tex_math_dollars+raw_tex",
        "--to=docx",
        f"--resource-path={HERE}{os.pathsep}{PROJECT}",
        f"--output={destination}",
    ]
    run(*command, cwd=HERE)


def localize_contents_bookmarks(document: Document) -> int:
    """Collapse contents targets onto their heading paragraphs for PDF conversion."""
    body = document.element.body
    paragraphs = list(body.iter(qn("w:p")))
    starts = {
        node.get(qn("w:name")): node
        for node in body.iter(qn("w:bookmarkStart"))
        if node.get(qn("w:name"))
    }
    ends = {
        node.get(qn("w:id")): node
        for node in body.iter(qn("w:bookmarkEnd"))
        if node.get(qn("w:id"))
    }

    localized = 0
    for reference in contents_references():
        start = starts.get(reference.anchor)
        if start is None:
            raise ValueError(f"DOCX lacks contents bookmark target: {reference.anchor}")
        end = ends.get(start.get(qn("w:id")))
        if end is None:
            raise ValueError(f"DOCX bookmark has no end marker: {reference.anchor}")

        heading = None
        candidates = []
        if start.getparent() is body:
            candidate = start.getnext()
            while candidate is not None:
                if candidate.tag == qn("w:p"):
                    candidates.append(candidate)
                candidate = candidate.getnext()
        else:
            start_paragraph = start
            while start_paragraph is not None and start_paragraph.tag != qn("w:p"):
                start_paragraph = start_paragraph.getparent()
            if start_paragraph is None:
                raise ValueError(f"DOCX bookmark has no paragraph context: {reference.anchor}")
            start_index = paragraphs.index(start_paragraph)
            candidates = paragraphs[start_index + 1 :]
        for candidate in candidates:
            style = candidate.find(f"{qn('w:pPr')}/{qn('w:pStyle')}")
            style_name = style.get(qn("w:val")) if style is not None else ""
            if style_name in {"Heading1", "Heading2"}:
                heading = candidate
                break
        if heading is None:
            raise ValueError(f"DOCX bookmark is not followed by a heading: {reference.anchor}")

        start.getparent().remove(start)
        end.getparent().remove(end)
        insert_at = 1 if heading.find(qn("w:pPr")) is not None else 0
        heading.insert(insert_at, start)
        heading.append(end)
        localized += 1
    return localized


def style_document(
    raw_docx: Path,
    destination: Path,
    *,
    audience: str,
    single_reveal: bool = False,
    document_title: str | None = None,
) -> None:
    document = Document(raw_docx)
    if audience == "instructor":
        title_text = "Teaching Cø / N*: A Graduate Instructor's Manual"
        subject_text = (
            "Audited lecture notes, theory placement, collaborative exercises, validation arguments, "
            "assessment materials, and instructor keys"
        )
        header_text = "TEACHING Cø / N*  |  INSTRUCTOR EDITION - KEYS INCLUDED"
    elif audience == "student":
        title_text = "Cø / N* Student Base Resource Pack"
        subject_text = "Key-free and reveal-safe collaborative materials for the Cø / N* graduate seminar"
        header_text = "Cø / N*  |  STUDENT BASE PACK - NO KEYS OR REVEALS"
    elif audience == "reveal":
        title_text = "Cø / N* Instructor-Controlled Student Reveal Sheets"
        subject_text = "Key-free controlled evidence releases for the Cø / N* graduate seminar"
        header_text = "Cø / N*  |  INSTRUCTOR-CONTROLLED REVEAL - NO KEY"
    else:
        raise ValueError(f"Unknown document audience: {audience}")

    document.core_properties.title = document_title or title_text
    document.core_properties.subject = subject_text
    document.core_properties.author = "Phil Stilwell"
    document.core_properties.last_modified_by = "Cø / N* teacher-manual build"
    document.core_properties.comments = f"Audited edition {EDITION}; corpus baseline: {CORPUS_BASELINE}"
    document.core_properties.keywords = (
        "consciousness; phenomenal presence; N*; graduate course; integration; availability; "
        "recurrence; system boundaries; indeterminacy; collaborative learning"
    )

    settings = document.settings._element
    theme = settings.find(qn("w:themeFontLang"))
    if theme is None:
        theme = OxmlElement("w:themeFontLang")
        settings.append(theme)
    theme.set(qn("w:val"), "en-US")

    for section in document.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.78)
        section.right_margin = Inches(0.78)
        section.header_distance = Inches(0.28)
        section.footer_distance = Inches(0.3)
        section.different_first_page_header_footer = not single_reveal

        header = section.header.paragraphs[0]
        header.text = header_text
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for item in header.runs:
            set_run_font(item, DISPLAY_FONT, 8)
            item.font.bold = True
            item.font.color.rgb = RGBColor(46, 111, 142)

        footer = section.footer.paragraphs[0]
        add_page_field(footer)
        for item in footer.runs:
            set_run_font(item, DISPLAY_FONT, 8)
            item.font.color.rgb = RGBColor(90, 99, 107)

        if not single_reveal:
            section.first_page_header.paragraphs[0].text = ""
            section.first_page_footer.paragraphs[0].text = ""

    if audience == "instructor":
        localized = localize_contents_bookmarks(document)
        if localized != len(contents_references()):
            raise ValueError("Not every instructor contents bookmark was localized")

    for name in ("Normal", "Body Text", "First Paragraph", "Compact"):
        if name in document.styles:
            style = document.styles[name]
            set_font(style, BODY_FONT, 10.25)
            style.font.color.rgb = TEXT
            # The student handout carries dense case materials but should still
            # keep each session's accessibility note on the session page.
            style.paragraph_format.line_spacing = 1.06 if audience == "student" else 1.08
            style.paragraph_format.space_after = Pt(4.5 if audience == "student" else 5.5)
            style.paragraph_format.widow_control = True

    title = document.styles["Title"]
    set_font(title, DISPLAY_FONT, 27, bold=True)
    title.font.color.rgb = RGBColor(27, 52, 76)
    title.paragraph_format.space_before = Pt(92)
    title.paragraph_format.space_after = Pt(11)

    if "Subtitle" in document.styles:
        subtitle = document.styles["Subtitle"]
        set_font(subtitle, BODY_FONT, 13, italic=True)
        subtitle.font.color.rgb = RGBColor(65, 86, 102)
        subtitle.paragraph_format.space_after = Pt(26)

    if "Author" in document.styles:
        author = document.styles["Author"]
        set_font(author, DISPLAY_FONT, 10.5, bold=True)
        author.font.color.rgb = RGBColor(46, 111, 142)
        author.paragraph_format.space_after = Pt(5)

    if "Date" in document.styles:
        date = document.styles["Date"]
        set_font(date, BODY_FONT, 9.5)
        date.font.color.rgb = RGBColor(90, 99, 107)

    heading_specs = {
        "Heading 1": (DISPLAY_FONT, 18, 18, 8),
        "Heading 2": (DISPLAY_FONT, 14, 14, 6),
        "Heading 3": (DISPLAY_FONT, 11, 10, 4),
        "Heading 4": (DISPLAY_FONT, 10, 8, 3),
    }
    for name, (font, size, before, after) in heading_specs.items():
        if name not in document.styles:
            continue
        style = document.styles[name]
        set_font(style, font, size, bold=True)
        style.font.color.rgb = RGBColor(27, 52, 76) if name != "Heading 3" else RGBColor(46, 111, 142)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        if name == "Heading 1":
            style.paragraph_format.page_break_before = True

    if "Block Text" in document.styles:
        block = document.styles["Block Text"]
        set_font(block, BODY_FONT, 10.25, italic=True)
        block.font.color.rgb = RGBColor(27, 52, 76)
        block.paragraph_format.left_indent = Inches(0.25)
        block.paragraph_format.right_indent = Inches(0.18)
        block.paragraph_format.space_before = Pt(7)
        block.paragraph_format.space_after = Pt(7)

    worksheet_intro_pending = False
    compact_session_12 = False
    compact_appendix_f = False
    compact_resource_session_4 = False
    compact_student_session_10 = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name == "Heading 2":
            if audience == "instructor":
                if text.startswith("Session 12."):
                    compact_session_12 = True
                elif compact_session_12:
                    compact_session_12 = False
                if text.startswith("Appendix F."):
                    compact_appendix_f = True
                elif compact_appendix_f:
                    compact_appendix_f = False
                if text.startswith("Session 4 Resource."):
                    compact_resource_session_4 = True
                elif compact_resource_session_4:
                    compact_resource_session_4 = False
            if audience == "student":
                if text.startswith("Session 10 Resource."):
                    compact_student_session_10 = True
                elif compact_student_session_10:
                    compact_student_session_10 = False
        if (
            paragraph.style.name not in {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}
            and (
                compact_session_12
                or compact_appendix_f
                or compact_resource_session_4
                or compact_student_session_10
            )
        ):
            paragraph.paragraph_format.line_spacing = (
                1.02 if compact_resource_session_4 or compact_student_session_10 else 1.04
            )
            paragraph.paragraph_format.space_after = Pt(
                3.5 if compact_resource_session_4 or compact_student_session_10 else 4
            )
        if (
            audience == "instructor"
            and paragraph.style.name == "Heading 3"
            and text.startswith("Worksheet ")
        ):
            if not text.startswith("Worksheet 1."):
                paragraph.paragraph_format.page_break_before = True
            worksheet_intro_pending = True
        elif worksheet_intro_pending and text:
            paragraph.paragraph_format.keep_with_next = True
            worksheet_intro_pending = False
        if (
            audience == "reveal"
            and not single_reveal
            and paragraph.style.name == "Heading 2"
            and text.startswith("Reveal Sheet -")
        ):
            paragraph.paragraph_format.page_break_before = True
        session_heading = re.match(r"^Session (\d+)(?:\.|\s)", text)
        if paragraph.style.name == "Heading 2" and (
            (session_heading is not None and int(session_heading.group(1)) != 1)
            or (text.startswith("Appendix ") and not text.startswith("Appendix A."))
        ):
            paragraph.paragraph_format.page_break_before = True
        if paragraph.style.name == "Heading 2" and text.startswith("6. The 14-session sequence"):
            paragraph.paragraph_format.page_break_before = True
        if paragraph.style.name == "Heading 3" and text == "Capstone rubric":
            paragraph.paragraph_format.page_break_before = True
        if text == "Preface" and paragraph.style.name == "Heading 1":
            paragraph.paragraph_format.page_break_before = True
        if paragraph.style.name == "Block Text":
            set_paragraph_shading(paragraph, PALE_BLUE)
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)

    for shape in document.inline_shapes:
        if shape.width > Inches(6.75):
            ratio = shape.height / shape.width
            shape.width = Inches(6.75)
            shape.height = int(shape.width * ratio)

    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        if not table.rows:
            continue
        header_signature, column_weights, total_width, table_font_size = table_layout_profile(table)
        compact_student_canvas = audience == "student" and header_signature == ("Field", "Team entry")
        if compact_student_canvas:
            table_font_size = 8.5
        set_fixed_table_columns(table, fixed_column_twips(column_weights, total_width))
        # Formula-only headers must inherit the navy-band typography. Dense
        # numerical tables also need ordinary text metrics so adjacent intervals
        # retain a visible gutter in both Word and LibreOffice output.
        if header_signature in DENSE_TABLE_FONT_SIZES:
            flatten_office_math(table._tbl)
        else:
            flatten_office_math(table.rows[0]._tr)
        is_contents = header_signature == ("Part or session", "Focus", "PDF page")
        centered_columns = set(CENTERED_TABLE_COLUMNS.get(header_signature, ()))
        worksheet_row_heights = {
            ("Field", "Entry"): 0.48,
            ("Question", "Cø / N*", "Rival theory"): 0.62,
            ("Step", "What to include", "Team record"): 0.42,
            ("Question", "Entry"): 0.48,
            ("Condition", "Standard to inspect", "Decision and case-specific evidence"): 0.72,
            (
                "Evidence round",
                "Provisional output",
                "What changed?",
                "System change or license change?",
                "Next discriminating evidence",
            ): 0.88,
            (
                "Criticism",
                "Source role or team",
                "Fatal, strengthening, or optional",
                "Accepted?",
                "Protocol change or reason for rejection",
            ): 1.18,
        }
        worksheet_row_height = (
            worksheet_row_heights.get(header_signature) if audience == "instructor" else None
        )
        repeat_table_header(table.rows[0])
        for row_index, row in enumerate(table.rows):
            prevent_row_split(row)
            if worksheet_row_height is not None and row_index > 0:
                row.height = Inches(worksheet_row_height)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for column_index, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if is_contents:
                    set_cell_margins(cell, top=28, bottom=28, start=55, end=55)
                elif header_signature == ("Term or symbol", "Teaching definition", "Guardrail"):
                    set_cell_margins(cell, top=45, bottom=45, start=75, end=75)
                elif compact_student_canvas:
                    set_cell_margins(cell, top=32, bottom=32, start=70, end=70)
                elif header_signature in DENSE_TABLE_FONT_SIZES:
                    set_cell_margins(cell, top=55, bottom=55, start=40, end=40)
                else:
                    set_cell_margins(cell)
                fill = NAVY if row_index == 0 else (PALE_GRAY if row_index % 2 == 0 else WHITE)
                set_shading(cell._tc, fill)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(
                        0 if is_contents or compact_student_canvas else 2
                    )
                    paragraph.paragraph_format.line_spacing = (
                        0.92 if is_contents else (0.94 if compact_student_canvas else 1.0)
                    )
                    if column_index in centered_columns:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_math_run_format(
                        paragraph,
                        table_font_size,
                        WHITE if row_index == 0 else str(TEXT),
                    )
                    for item in paragraph.runs:
                        set_run_font(item, DISPLAY_FONT, 7.8 if is_contents else table_font_size)
                        if row_index == 0:
                            item.font.bold = True
                            item.font.color.rgb = RGBColor(255, 255, 255)
                        else:
                            item.font.color.rgb = TEXT

    for paragraph in document.paragraphs:
        for hyperlink in paragraph._p.xpath(".//w:hyperlink"):
            for item in hyperlink.xpath(".//w:r"):
                r_pr = item.get_or_add_rPr()
                color = r_pr.find(qn("w:color"))
                if color is None:
                    color = OxmlElement("w:color")
                    r_pr.append(color)
                color.set(qn("w:val"), BLUE)

    document.save(destination)


def convert_pdf(soffice: str, docx_path: Path, work_dir: Path) -> Path:
    profile_dir = Path(tempfile.mkdtemp(prefix="c0-manual-lo-"))
    try:
        result = run(
            soffice,
            f"-env:UserInstallation={profile_dir.as_uri()}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(work_dir),
            str(docx_path),
        )
    finally:
        shutil.rmtree(profile_dir, ignore_errors=True)
    generated = work_dir / f"{docx_path.stem}.pdf"
    if not generated.exists():
        raise FileNotFoundError(
            f"LibreOffice did not create {generated}. stdout={result.stdout!r}; stderr={result.stderr!r}"
        )
    return generated


def flatten_outline(items) -> list[str]:
    titles: list[str] = []
    for item in items:
        if isinstance(item, list):
            titles.extend(flatten_outline(item))
        else:
            title = getattr(item, "title", None)
            if title:
                titles.append(str(title))
    return titles


def flatten_outline_pages(reader: PdfReader, items=None) -> list[tuple[str, int]]:
    """Return all PDF outline titles with their zero-based destination pages."""
    if items is None:
        items = reader.outline
    destinations: list[tuple[str, int]] = []
    for item in items:
        if isinstance(item, list):
            destinations.extend(flatten_outline_pages(reader, item))
            continue
        title = getattr(item, "title", None)
        if not title:
            continue
        page = reader.get_destination_page_number(item)
        if page < 0:
            raise ValueError(f"PDF outline item has no valid destination page: {title}")
        destinations.append((str(title), page))
    return destinations


def outline_page_for_label(outline: list[tuple[str, int]], label: str) -> int:
    matches = [
        page
        for title, page in outline
        if title == label or title.startswith(f"{label}.")
    ]
    if len(matches) != 1:
        raise ValueError(f"Expected one outline destination for {label!r}; found {matches}")
    return matches[0]


def annotation_destination_page(reader: PdfReader, annotation) -> int | None:
    annotation = annotation.get_object()
    destination = annotation.get("/Dest")
    if destination is None:
        action = annotation.get("/A")
        if action is not None and action.get_object().get("/S") == "/GoTo":
            destination = action.get_object().get("/D")
    if not isinstance(destination, list) or not destination:
        return None
    target = destination[0]
    target_id = getattr(target, "idnum", None)
    if target_id is None:
        return None
    page_by_id = {
        page.indirect_reference.idnum: index
        for index, page in enumerate(reader.pages)
        if page.indirect_reference is not None
    }
    return page_by_id.get(target_id)


def validate_contents_pdf(reader: PdfReader) -> dict[str, object]:
    """Validate printed page numbers and live PDF links for all contents rows."""
    references = contents_references()
    outline = flatten_outline_pages(reader)
    contents_page = outline_page_for_label(outline, "Contents")
    page = reader.pages[contents_page]
    chunks: list[tuple[str, float, float]] = []

    def collect(text, _cm, tm, _font, _size) -> None:
        clean = text.strip()
        if clean:
            chunks.append((clean, float(tm[4]), float(tm[5])))

    page.extract_text(visitor_text=collect)
    annotations = [
        annotation
        for annotation in (page.get("/Annots") or [])
        if annotation.get_object().get("/Subtype") == "/Link"
    ]
    verified: dict[str, int] = {}
    for reference in references:
        expected_page = outline_page_for_label(outline, reference.label)
        if reference.printed_page != expected_page + 1:
            raise ValueError(
                f"Printed contents page for {reference.label!r} is {reference.printed_page}; "
                f"outline destination is {expected_page + 1}"
            )
        positions = [
            (x, y)
            for text, x, y in chunks
            if text == reference.label
            or (
                reference.label.startswith("Part ")
                and len(text) >= 20
                and reference.label.startswith(text)
            )
        ]
        if len(positions) != 1:
            raise ValueError(f"Could not locate one contents label on the PDF page: {reference.label}")
        x, y = positions[0]
        matching_links = []
        for annotation in annotations:
            rect = annotation.get_object().get("/Rect")
            if rect is None:
                continue
            x1, y1, x2, y2 = (float(value) for value in rect)
            if x1 - 1 <= x <= x2 + 1 and y1 - 1 <= y <= y2 + 1:
                matching_links.append(annotation_destination_page(reader, annotation))
        if expected_page not in matching_links:
            raise ValueError(
                f"Contents link for {reference.label!r} does not resolve to page {expected_page + 1}; "
                f"found={matching_links}"
            )
        verified[reference.anchor] = expected_page + 1
    return {
        "printed_references": len(references),
        "live_links_verified": len(verified),
        "destinations": verified,
    }


def validate_worksheet_pages(page_text: list[str]) -> dict[str, int]:
    pages: dict[str, int] = {}
    for number in range(1, 10):
        token = f"Worksheet {number}."
        matches = [index + 1 for index, text in enumerate(page_text) if token in text]
        if len(matches) != 1:
            raise ValueError(f"Worksheet {number} occurs on PDF pages {matches}; expected one")
        pages[str(number)] = matches[0]
    if len(set(pages.values())) != 9:
        raise ValueError(f"Worksheets do not occupy nine distinct pages: {pages}")
    return pages


def pdf_font_report(pdf_path: Path) -> dict[str, object]:
    pdffonts = shutil.which("pdffonts")
    if not pdffonts:
        return {"checked": False, "reason": "pdffonts not available"}
    lines = tool_output(pdffonts, str(pdf_path)).splitlines()[2:]
    records = [line.split() for line in lines if line.strip()]
    malformed = [parts for parts in records if len(parts) < 8]
    if malformed:
        raise ValueError(f"Could not parse pdffonts output for {pdf_path}: {malformed}")
    embedded = all(parts[-5] == "yes" for parts in records)
    unicode_mapped = all(parts[-3] == "yes" for parts in records)
    families = sorted({parts[0].split("+", 1)[-1] for parts in records})
    if not embedded:
        raise ValueError(f"PDF contains an unembedded font: {pdf_path}")
    normalized = {re.sub(r"[^a-z0-9]", "", family.lower()) for family in families}
    for requested in (BODY_FONT, DISPLAY_FONT):
        needle = re.sub(r"[^a-z0-9]", "", requested.lower())
        if not any(needle in family for family in normalized):
            raise ValueError(f"Requested font {requested!r} is absent from {pdf_path}; found {families}")
    return {
        "checked": True,
        "tool": Path(pdffonts).name,
        "font_programs": len(records),
        "families": families,
        "all_embedded": embedded,
        "all_unicode_mapped": unicode_mapped,
    }


def validate_artifact(
    docx_path: Path,
    pdf_path: Path,
    *,
    audience: str,
    reveal_titles: list[str] | None = None,
) -> dict[str, object]:
    minimum_bytes = 10_000 if audience == "reveal" else 25_000
    if docx_path.stat().st_size < minimum_bytes:
        raise ValueError(f"DOCX is unexpectedly small: {docx_path}")
    if pdf_path.stat().st_size < minimum_bytes:
        raise ValueError(f"PDF is unexpectedly small: {pdf_path}")

    with ZipFile(docx_path) as archive:
        bad_member = archive.testzip()
        if bad_member:
            raise ValueError(f"Corrupt DOCX member: {bad_member}")
        names = archive.namelist()
        xml = archive.read("word/document.xml").decode("utf-8")
        media_count = sum(name.startswith("word/media/") for name in names)
        internal_links = xml.count("w:anchor=")
        bookmarks = xml.count("w:bookmarkStart")
        alt_text_count = xml.count("descr=")

    doc = Document(docx_path)
    table_layout_report = validate_docx_table_layouts(doc)
    doc_text = "\n".join(node.text or "" for node in doc.element.body.iter(qn("w:t")))
    reader = PdfReader(str(pdf_path))
    page_text = [(page.extract_text() or "") for page in reader.pages]
    text = "\n".join(page_text)
    outline_titles = flatten_outline(reader.outline)
    root = reader.trailer["/Root"]
    mark_info = root.get("/MarkInfo")
    marked = bool(mark_info and mark_info.get_object().get("/Marked"))
    language = str(root.get("/Lang") or "")
    tagged = "/StructTreeRoot" in root and marked
    if not tagged or language != "en-US":
        raise ValueError(
            f"PDF accessibility catalog is incomplete: tagged={tagged}, language={language!r}"
        )

    searchable = f"{doc_text}\n{text}"
    normalized_doc_text = re.sub(r"\s+", " ", doc_text)
    normalized_pdf_text = re.sub(r"\s+", " ", text)
    if audience == "instructor":
        required = ("Preface", "Session 14", "Appendix H", "Closing instructor note")
        minimum_pages = 100
        if doc_text.count("Instructor Key") != 14:
            raise ValueError("Instructor DOCX does not contain exactly 14 answer-key headings")
        if text.count("Instructor Key") != 14:
            raise ValueError("Instructor edition does not contain exactly 14 answer-key headings")
        if media_count != 5:
            raise ValueError(f"Instructor edition should contain 5 figures; found {media_count}")
        if internal_links < 20:
            raise ValueError(f"Instructor edition should contain internal navigation links; found {internal_links}")
        contents_report = validate_contents_pdf(reader)
        worksheet_pages = validate_worksheet_pages(page_text)
    elif audience == "student":
        required = ("Student Use Guide", "Session 1 Resource", "Session 14 Resource")
        minimum_pages = 20
        forbidden = ("Instructor Key", "Instructor Preparation Crosswalk", "Reveal Sheet -", "KEYS INCLUDED")
        leaked = [token for token in forbidden if token in searchable]
        if leaked:
            raise ValueError(f"Student base pack contains restricted content: {leaked}")
    elif audience == "reveal":
        required = (
            "Reveal-Sheet Use Guide",
            "Reveal Sheet - Session 4",
            "Reveal Sheet - Session 11",
            "Reveal Sheet - Session 14",
        )
        minimum_pages = 10
        forbidden = ("Instructor Key", "Instructor Preparation Crosswalk", "KEYS INCLUDED")
        leaked = [token for token in forbidden if token in searchable]
        if leaked:
            raise ValueError(f"Controlled-reveal pack contains instructor-only content: {leaked}")
        if reveal_titles is None:
            raise ValueError("Controlled-reveal validation requires the audited title inventory")
        doc_headings = [
            paragraph.text.strip()
            for paragraph in doc.paragraphs
            if paragraph.style.name == "Heading 2"
            and paragraph.text.strip().startswith("Reveal Sheet - Session ")
        ]
        pdf_headings = [
            title for title in outline_titles if title.startswith("Reveal Sheet - Session ")
        ]
        if doc_headings != reveal_titles:
            raise ValueError("Controlled-reveal DOCX heading inventory or order is incomplete")
        if pdf_headings != reveal_titles:
            raise ValueError("Controlled-reveal PDF outline inventory or order is incomplete")
        for title in reveal_titles:
            if normalized_doc_text.count(title) != 1:
                raise ValueError(f"Controlled-reveal DOCX does not contain exactly one {title!r}")
            if normalized_pdf_text.count(title) != 1:
                raise ValueError(f"Controlled-reveal PDF does not contain exactly one {title!r}")
    else:
        raise ValueError(f"Unknown artifact audience: {audience}")

    missing = [token for token in required if token not in doc_text]
    if missing:
        raise ValueError(f"Required document text is missing: {missing}")
    if len(reader.pages) < minimum_pages:
        raise ValueError(f"PDF has only {len(reader.pages)} pages; expected at least {minimum_pages}")
    if not outline_titles:
        raise ValueError("PDF contains no navigation outline")

    artifacts = ("N*,*", "N model*", "{#appendix", "{#session", "$Pi_c$")
    visible_artifacts = [token for token in artifacts if token in text or token in doc_text]
    if visible_artifacts:
        raise ValueError(f"Visible source artifacts remain: {visible_artifacts}")

    near_blank_pages = [index + 1 for index, value in enumerate(page_text) if len(value.strip()) < 12]
    if near_blank_pages:
        raise ValueError(f"PDF contains near-blank pages: {near_blank_pages}")
    sparse_pages = []
    if audience in {"instructor", "student"}:
        sparse_pages = [
            index + 1
            for index, value in enumerate(page_text)
            if index > 0 and len(value.strip()) < 400
        ]
        if sparse_pages:
            raise ValueError(f"PDF contains materially sparse continuation pages: {sparse_pages}")
    return {
        "docx_bytes": docx_path.stat().st_size,
        "pdf_bytes": pdf_path.stat().st_size,
        "pages": len(reader.pages),
        "docx_paragraphs": len(doc.paragraphs),
        "docx_tables": len(doc.tables),
        "table_layout": table_layout_report,
        "media_count": media_count,
        "alt_text_count": alt_text_count,
        "internal_links": internal_links,
        "bookmarks": bookmarks,
        "pdf_outline_entries": len(outline_titles),
        "near_blank_pages": near_blank_pages,
        "sparse_continuation_pages": sparse_pages,
        "tagged": tagged,
        "language": language,
        **({"contents": contents_report, "worksheet_pages": worksheet_pages} if audience == "instructor" else {}),
        "font_report": pdf_font_report(pdf_path),
    }


def build_individual_reveal_pdfs(
    pandoc: str,
    soffice: str,
    reveal_blocks: list[AudienceBlock],
    work_dir: Path,
    destination_dir: Path,
) -> list[dict[str, object]]:
    """Build each reveal independently so tagging and language are preserved."""
    specs = [reveal_spec(block) for block in reveal_blocks]
    destination_dir.mkdir(parents=True, exist_ok=False)
    reveal_work = work_dir / "individual-reveal-work"
    reveal_work.mkdir(parents=True, exist_ok=False)
    metadata: list[dict[str, object]] = []
    all_titles = [spec.title for spec in specs]
    for block, spec in zip(reveal_blocks, specs, strict=True):
        filename = f"c0-n-star-{spec.block_id}-reveal.pdf"
        output = destination_dir / filename
        source = reveal_work / f"{spec.block_id}.md"
        source.write_text(
            "---\nlang: en-US\n---\n\n"
            + re.sub(r"^### Reveal Sheet", "## Reveal Sheet", block.markdown, count=1)
            + "\n",
            encoding="utf-8",
        )
        raw_docx = reveal_work / f"{spec.block_id}-raw.docx"
        styled_docx = reveal_work / f"{Path(filename).stem}.docx"
        pandoc_docx(pandoc, [source], raw_docx)
        style_document(
            raw_docx,
            styled_docx,
            audience="reveal",
            single_reveal=True,
            document_title=spec.title,
        )
        generated_pdf = convert_pdf(soffice, styled_docx, reveal_work)
        os.replace(generated_pdf, output)

        reveal_reader = PdfReader(str(output))
        if len(reveal_reader.pages) != 1:
            raise ValueError(
                f"Reveal sheet must be exactly one page: {spec.block_id} has {len(reveal_reader.pages)}"
            )
        page_text = [(page.extract_text() or "") for page in reveal_reader.pages]
        normalized = re.sub(r"\s+", " ", "\n".join(page_text))
        if normalized.count(spec.title) != 1:
            raise ValueError(f"Individual reveal lacks its unique heading: {spec.block_id}")
        leaked_titles = [
            title for title in all_titles if title != spec.title and title in normalized
        ]
        if leaked_titles:
            raise ValueError(
                f"Individual reveal {spec.block_id} contains another release heading: {leaked_titles}"
            )
        if "Instructor Key" in normalized or "Instructor Preparation Crosswalk" in normalized:
            raise ValueError(f"Individual reveal contains instructor-only content: {spec.block_id}")
        near_blank = [number + 1 for number, value in enumerate(page_text) if len(value.strip()) < 12]
        if near_blank:
            raise ValueError(f"Individual reveal contains near-blank pages: {spec.block_id} {near_blank}")
        if output.stat().st_size < 5_000:
            raise ValueError(f"Individual reveal PDF is unexpectedly small: {output}")
        root = reveal_reader.trailer["/Root"]
        mark_info = root.get("/MarkInfo")
        marked = bool(mark_info and mark_info.get_object().get("/Marked"))
        tagged = "/StructTreeRoot" in root and marked
        language = str(root.get("/Lang") or "")
        if not tagged or language != "en-US":
            raise ValueError(
                f"Individual reveal accessibility metadata failed: {spec.block_id}; "
                f"tagged={tagged}, language={language!r}"
            )
        document_info = reveal_reader.metadata
        if document_info is None or document_info.title != spec.title:
            raise ValueError(f"Individual reveal title metadata failed: {spec.block_id}")
        if document_info.author != "Phil Stilwell" or not document_info.subject:
            raise ValueError(f"Individual reveal author/subject metadata failed: {spec.block_id}")
        outline_titles = flatten_outline(reveal_reader.outline)
        if outline_titles != [spec.title]:
            raise ValueError(f"Individual reveal outline failed: {spec.block_id} {outline_titles}")
        font_report = pdf_font_report(output)
        metadata.append(
            {
                "block_id": spec.block_id,
                "session": spec.session,
                "title": spec.title,
                "pages": 1,
                "filename": filename,
                "sha256": sha256(output),
                "tagged": tagged,
                "language": language,
                "outline_entries": len(outline_titles),
                "fonts_embedded": font_report.get("all_embedded"),
            }
        )
    return metadata


def remove_path(path: Path) -> None:
    if path.is_dir() and not path.is_symlink():
        shutil.rmtree(path)
    elif path.exists() or path.is_symlink():
        path.unlink()


def transactional_publish(
    file_pairs: list[tuple[Path, Path]],
    directory_pair: tuple[Path, Path],
    verifier: Callable[[], None],
) -> None:
    """Publish and verify the complete release set, rolling back on any failure."""
    staged_directory, destination_directory = directory_pair
    if not file_pairs or file_pairs[-1][1] != MANIFEST_OUT:
        raise ValueError("The manifest must be the final file pair in a release transaction")
    pairs = [*file_pairs[:-1], (staged_directory, destination_directory), file_pairs[-1]]
    backup_root = Path(tempfile.mkdtemp(prefix="release-backup-", dir=TMP_ROOT))
    backups: list[tuple[Path, Path]] = []
    installed: list[Path] = []
    try:
        for index, (staged, destination) in enumerate(pairs):
            if not staged.exists():
                raise FileNotFoundError(f"Staged release component is missing: {staged}")
            destination.parent.mkdir(parents=True, exist_ok=True)
            if destination.exists() or destination.is_symlink():
                backup = backup_root / f"{index:02d}-{destination.name}"
                os.replace(destination, backup)
                backups.append((backup, destination))
            os.replace(staged, destination)
            installed.append(destination)
        verifier()
    except Exception:
        for destination in reversed(installed):
            remove_path(destination)
        for backup, destination in reversed(backups):
            destination.parent.mkdir(parents=True, exist_ok=True)
            os.replace(backup, destination)
        raise
    finally:
        shutil.rmtree(backup_root, ignore_errors=True)


def verify_published_release(manifest: dict[str, object]) -> None:
    """Reject a mixed or stale published set after the transactional swap."""
    for relative, expected in manifest["sources"].items():
        path = PROJECT / relative
        if sha256(path) != expected:
            raise ValueError(f"Published manifest has a stale source hash: {relative}")

    for relative, expected in manifest["corpus_sources"].items():
        path = PROJECT / relative
        if sha256(path) != expected:
            raise ValueError(f"Published manifest has a stale corpus hash: {relative}")

    for output in manifest["outputs"].values():
        for relative, expected in output["files"].items():
            path = PROJECT / relative
            if not path.is_file() or sha256(path) != expected:
                raise ValueError(f"Published artifact hash mismatch: {relative}")
        for relative, expected in output["normalized_content_sha256"].items():
            path = PROJECT / relative
            if normalized_content_sha256(path) != expected:
                raise ValueError(f"Published normalized-content mismatch: {relative}")

    reveal_inventory = manifest["reveal_sheets"]
    expected_paths: set[Path] = set()
    for record in reveal_inventory.values():
        path = PROJECT / record["path"]
        expected_paths.add(path)
        if not path.is_file() or sha256(path) != record["sha256"]:
            raise ValueError(f"Published split-reveal hash mismatch: {record['path']}")
    published_paths = set((PDF_DIR / "reveals").glob("*.pdf"))
    if published_paths != expected_paths:
        raise ValueError("Published split-reveal directory does not match the manifest inventory")

    disk_manifest = json.loads(MANIFEST_OUT.read_text(encoding="utf-8"))
    if disk_manifest != manifest:
        raise ValueError("Published manifest bytes do not describe the verified release")


def build_release() -> None:
    if sys.version_info < (3, 10):
        raise RuntimeError("Python 3.10 or later is required")
    inputs = source_inputs()
    source_qa = validate_sources(inputs)

    pandoc = find_tool("pandoc", "C0_PANDOC")
    soffice = find_tool("soffice", "C0_SOFFICE")
    TMP_ROOT.mkdir(parents=True, exist_ok=True)
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    work_dir = Path(tempfile.mkdtemp(prefix="build-", dir=TMP_ROOT))

    try:
        blocks = parse_audience_blocks()
        reveal_blocks = [block for block in blocks if block.audience == "reveal"]
        reveal_titles = [reveal_spec(block).title for block in reveal_blocks]
        generated_sources: dict[str, Path] = {}
        provenance: dict[str, list[str]] = {}
        for label in ("instructor", "student", "reveal"):
            markdown, block_ids = audience_markdown(label, blocks)
            generated_source = work_dir / f"{label}-resource-pack.md"
            generated_source.write_text(markdown, encoding="utf-8")
            generated_sources[label] = generated_source
            provenance[label] = block_ids

        staged: dict[str, tuple[Path, Path]] = {}
        qa: dict[str, dict[str, object]] = {}
        specifications = (
            ("instructor", [MANUAL_SOURCE, generated_sources["instructor"]], INSTRUCTOR_STEM),
            ("student", [generated_sources["student"]], STUDENT_STEM),
            ("reveal", [generated_sources["reveal"]], REVEAL_STEM),
        )
        for label, sources, stem in specifications:
            raw_docx = work_dir / f"{stem}-raw.docx"
            styled_docx = work_dir / f"{stem}.docx"
            pandoc_docx(pandoc, sources, raw_docx)
            style_document(raw_docx, styled_docx, audience=label)
            staged_pdf = convert_pdf(soffice, styled_docx, work_dir)
            qa[label] = validate_artifact(
                styled_docx,
                staged_pdf,
                audience=label,
                reveal_titles=reveal_titles if label == "reveal" else None,
            )
            staged[label] = (styled_docx, staged_pdf)

        staged_reveal_dir = work_dir / "reveal-sheets"
        reveal_sheet_metadata = build_individual_reveal_pdfs(
            pandoc, soffice, reveal_blocks, work_dir, staged_reveal_dir
        )
        qa["reveal_sheets"] = {
            "count": len(reveal_sheet_metadata),
            "total_pages": sum(record["pages"] for record in reveal_sheet_metadata),
            "every_sheet_one_page": all(record["pages"] == 1 for record in reveal_sheet_metadata),
            "every_heading_verified": True,
            "every_sheet_key_free": True,
            "every_sheet_tagged": all(record["tagged"] for record in reveal_sheet_metadata),
            "every_sheet_language_en_us": all(
                record["language"] == "en-US" for record in reveal_sheet_metadata
            ),
            "every_sheet_fonts_embedded": all(
                record["fonts_embedded"] for record in reveal_sheet_metadata
            ),
            "every_sheet_outline_verified": all(
                record["outline_entries"] == 1 for record in reveal_sheet_metadata
            ),
        }

        final_paths = {
            "instructor": (DOCX_DIR / f"{INSTRUCTOR_STEM}.docx", PDF_DIR / f"{INSTRUCTOR_STEM}.pdf"),
            "student": (DOCX_DIR / f"{STUDENT_STEM}.docx", PDF_DIR / f"{STUDENT_STEM}.pdf"),
            "reveal": (DOCX_DIR / f"{REVEAL_STEM}.docx", PDF_DIR / f"{REVEAL_STEM}.pdf"),
        }
        manifest = {
            "schema_version": 2,
            "edition": EDITION,
            "generated_at_utc": datetime.now(timezone.utc).isoformat(),
            "corpus_baseline": CORPUS_BASELINE,
            "corpus_sources": {
                str(path.relative_to(PROJECT)): sha256(path) for path in CORPUS_SOURCES
            },
            "fonts": {"body": BODY_FONT, "display": DISPLAY_FONT},
            "font_files": font_file_report(soffice),
            "tools": {
                "python": sys.version.split()[0],
                "python_docx": docx.__version__,
                "pypdf": pypdf.__version__,
                "pandoc": tool_output(pandoc, "--version").splitlines()[0],
                "libreoffice": tool_output(soffice, "--version").splitlines()[0],
                "platform": platform.platform(),
                "machine": platform.machine(),
            },
            "reproducibility": {
                "scope": "semantic-and-layout reproducibility; container timestamps, PDF IDs, and release timestamps may vary",
                "normalized_content_digests_exclude": [
                    "DOCX core timestamps",
                    "PDF metadata dates and document IDs",
                ],
            },
            "sources": {str(path.relative_to(PROJECT)): sha256(path) for path in inputs},
            "outputs": {},
            "qa": {"sources": source_qa, **qa},
        }
        for label, final in final_paths.items():
            staged_paths = staged[label]
            manifest["outputs"][label] = {
                "files": {
                    str(destination.relative_to(PROJECT)): sha256(staged_path)
                    for staged_path, destination in zip(staged_paths, final)
                },
                "normalized_content_sha256": {
                    str(destination.relative_to(PROJECT)): normalized_content_sha256(staged_path)
                    for staged_path, destination in zip(staged_paths, final)
                },
                "included_block_ids": provenance[label],
                "generated_markdown_sha256": sha256(generated_sources[label]),
            }
        manifest["reveal_sheets"] = {
            record["block_id"]: {
                "session": record["session"],
                "title": record["title"],
                "pages": record["pages"],
                "path": str((PDF_DIR / "reveals" / record["filename"]).relative_to(PROJECT)),
                "sha256": record["sha256"],
                "tagged": record["tagged"],
                "language": record["language"],
                "outline_entries": record["outline_entries"],
                "fonts_embedded": record["fonts_embedded"],
            }
            for record in reveal_sheet_metadata
        }
        staged_manifest = work_dir / MANIFEST_OUT.name
        staged_manifest.write_text(json.dumps(manifest, indent=2, sort_keys=True) + "\n", encoding="utf-8")

        file_pairs: list[tuple[Path, Path]] = []
        for label in ("instructor", "student", "reveal"):
            file_pairs.extend(zip(staged[label], final_paths[label]))
        file_pairs.append((staged_manifest, MANIFEST_OUT))
        transactional_publish(
            file_pairs,
            (staged_reveal_dir, PDF_DIR / "reveals"),
            lambda: verify_published_release(manifest),
        )

        for label in ("instructor", "student", "reveal"):
            print(final_paths[label][0])
            print(final_paths[label][1])
        print(PDF_DIR / "reveals")
        print(MANIFEST_OUT)
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)


def build() -> None:
    """Serialize release builds so concurrent publishers cannot mix artifacts."""
    lock_dir = PROJECT / "tmp" / "docs"
    lock_dir.mkdir(parents=True, exist_ok=True)
    lock_path = lock_dir / "teachers-manual-build.lock"
    with lock_path.open("a+", encoding="utf-8") as lock_handle:
        try:
            fcntl.flock(lock_handle.fileno(), fcntl.LOCK_EX | fcntl.LOCK_NB)
        except BlockingIOError as exc:
            raise RuntimeError("Another teacher-manual build is already running") from exc
        try:
            build_release()
        finally:
            fcntl.flock(lock_handle.fileno(), fcntl.LOCK_UN)
    lock_path.unlink(missing_ok=True)


if __name__ == "__main__":
    try:
        build()
    except Exception as exc:
        print(f"Build failed: {exc}", file=sys.stderr)
        raise
