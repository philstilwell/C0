#!/usr/bin/env python3
"""Build polished DOCX and PDF artifacts from manuscript.md."""

from pathlib import Path
import subprocess

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = Path(__file__).with_name("manuscript.md")
TMP = ROOT / "tmp" / "docs"
DOC_OUT = ROOT / "output" / "doc" / "indeterminacy-as-a-scientific-result.docx"
PDF_OUT = ROOT / "output" / "pdf" / "indeterminacy-as-a-scientific-result.pdf"
RAW_DOCX = TMP / "indeterminacy-raw.docx"
PDF_SOURCE = TMP / "indeterminacy-pdf-source.md"


def run(*args: str) -> None:
    subprocess.run(args, check=True)


def set_font(style, name: str, size: int, bold=None, italic=None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_element = paragraph.add_run()._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run_element.extend([begin, instruction, end])


def style_docx() -> None:
    document = Document(RAW_DOCX)

    for section in document.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)
        header = section.header.paragraphs[0]
        header.text = "INDETERMINACY AS A SCIENTIFIC RESULT"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.runs[0].font.name = "Times New Roman"
        header.runs[0].font.size = Pt(9)
        header.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        add_page_number(section.footer.paragraphs[0])

    base_styles = ["Normal", "Body Text", "First Paragraph", "Compact"]
    for name in base_styles:
        if name in document.styles:
            style = document.styles[name]
            set_font(style, "Times New Roman", 11)
            style.paragraph_format.line_spacing = 1.15
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.widow_control = True

    title = document.styles["Title"]
    set_font(title, "Times New Roman", 22, bold=True)
    title.font.color.rgb = RGBColor(28, 45, 64)
    title.paragraph_format.space_before = Pt(80)
    title.paragraph_format.space_after = Pt(8)

    subtitle = document.styles["Subtitle"]
    set_font(subtitle, "Times New Roman", 14, italic=True)
    subtitle.font.color.rgb = RGBColor(66, 77, 88)
    subtitle.paragraph_format.space_after = Pt(28)

    if "Author" in document.styles:
        author = document.styles["Author"]
        set_font(author, "Times New Roman", 12)
        author.paragraph_format.space_after = Pt(8)
    if "Date" in document.styles:
        date = document.styles["Date"]
        set_font(date, "Times New Roman", 10)
        date.paragraph_format.line_spacing = 1.15

    heading_specs = {
        "Heading 1": (15, 12, 6),
        "Heading 2": (13, 14, 6),
        "Heading 3": (11, 10, 4),
    }
    for name, (size, before, after) in heading_specs.items():
        if name in document.styles:
            style = document.styles[name]
            set_font(style, "Times New Roman", size, bold=True)
            style.font.color.rgb = RGBColor(28, 45, 64)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True

    if "Block Text" in document.styles:
        block = document.styles["Block Text"]
        set_font(block, "Times New Roman", 10, italic=True)
        block.paragraph_format.left_indent = Inches(0.35)
        block.paragraph_format.right_indent = Inches(0.35)
        block.paragraph_format.space_before = Pt(6)
        block.paragraph_format.space_after = Pt(6)

    if not any(paragraph.text == "A visual preview:" for paragraph in document.paragraphs):
        for paragraph in document.paragraphs:
            if paragraph._p.xpath(".//w:drawing"):
                paragraph.insert_paragraph_before("A visual preview:")
                break

    for shape in document.inline_shapes:
        if shape.type is not None:
            shape.width = Inches(8.5)
            shape.height = Inches(8.5 * 1536 / 2752)

    for paragraph in document.paragraphs:
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Inches(-1.0)
            paragraph.paragraph_format.right_indent = Inches(-1.0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            break

    abstract_seen = False
    references_seen = False
    for paragraph in document.paragraphs:
        if paragraph.text == "A visual preview:":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(14)
            paragraph.paragraph_format.space_after = Pt(8)
            for run_item in paragraph.runs:
                run_item.font.name = "Times New Roman"
                run_item.font.size = Pt(11)
                run_item.font.bold = True
                run_item.font.color.rgb = RGBColor(28, 45, 64)
                run_item._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        elif paragraph.text == "Abstract":
            paragraph.paragraph_format.page_break_before = True
            abstract_seen = True
        elif paragraph.text == "References":
            paragraph.paragraph_format.page_break_before = True
            references_seen = True
            continue
        elif paragraph.text.startswith("Appendix ") and ":" in paragraph.text:
            references_seen = False
            paragraph.paragraph_format.page_break_before = True
        if references_seen and paragraph.text and paragraph.text != "References":
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.first_line_indent = Inches(-0.3)
            paragraph.paragraph_format.space_after = Pt(8)

    for table in document.tables:
        table.style = "Table"
        table.autofit = True
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run_item in paragraph.runs:
                        run_item.font.name = "Times New Roman"
                        run_item.font.size = Pt(9)
                        run_item._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                        if row_index == 0:
                            run_item.bold = True

    document.core_properties.title = "Indeterminacy as a Scientific Result"
    document.core_properties.subject = "A Four-Outcome Framework for Consciousness Attribution"
    document.core_properties.author = "Phil Stilwell"
    document.core_properties.keywords = (
        "consciousness attribution; indeterminacy; C-tests; diagnostic uncertainty"
    )
    DOC_OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(DOC_OUT)


def main() -> None:
    TMP.mkdir(parents=True, exist_ok=True)
    PDF_OUT.parent.mkdir(parents=True, exist_ok=True)
    preview_path = (SOURCE.parent / "visual-preview.png").as_posix()
    pdf_source = SOURCE.read_text(encoding="utf-8").replace(
        "![](visual-preview.png){width=100%}",
        f"\\noindent\\makebox[\\textwidth][c]{{\\includegraphics[width=\\paperwidth]{{{preview_path}}}}}",
    )
    PDF_SOURCE.write_text(pdf_source, encoding="utf-8")
    run(
        "pandoc",
        str(SOURCE),
        "--from=markdown+tex_math_dollars+tex_math_single_backslash",
        f"--resource-path={SOURCE.parent}",
        "--to=docx",
        f"--output={RAW_DOCX}",
    )
    style_docx()
    run(
        "pandoc",
        str(PDF_SOURCE),
        "--from=markdown+tex_math_dollars+tex_math_single_backslash",
        f"--resource-path={SOURCE.parent}",
        "--pdf-engine=xelatex",
        "-V",
        "papersize=letter",
        "-V",
        "geometry:margin=0.9in",
        "-V",
        "fontsize=11pt",
        "-V",
        "mainfont=Times New Roman",
        "-V",
        "sansfont=Arial",
        "-V",
        "header-includes=\\usepackage{graphicx}",
        "-V",
        "linestretch=1.15",
        "-V",
        "colorlinks=true",
        "-V",
        "linkcolor=NavyBlue",
        "-V",
        "urlcolor=NavyBlue",
        f"--output={PDF_OUT}",
    )


if __name__ == "__main__":
    main()
