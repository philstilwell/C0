#!/usr/bin/env python3
"""Build polished DOCX and PDF artifacts from manuscript.md."""

from pathlib import Path
import os
import subprocess

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(os.environ.get("PAPER_ROOT", Path(__file__).resolve().parents[2]))
SOURCE = Path(os.environ.get("PAPER_SOURCE", Path(__file__).with_name("manuscript.md")))
TMP = ROOT / "tmp" / "docs"
OUTPUT_STEM = os.environ.get("PAPER_OUTPUT_STEM", "where-is-the-conscious-subject")
DOC_OUT = ROOT / "output" / "doc" / f"{OUTPUT_STEM}.docx"
PDF_OUT = ROOT / "output" / "pdf" / f"{OUTPUT_STEM}.pdf"
RAW_DOCX = TMP / f"{OUTPUT_STEM}-raw.docx"
PDF_SOURCE = TMP / f"{OUTPUT_STEM}-pdf-source.md"
TABLE_FILTER = Path(os.environ.get("PAPER_TABLE_FILTER", SOURCE.parent / "table_layout.lua"))
PAPER_HEADER = os.environ.get("PAPER_HEADER", "WHERE IS THE CONSCIOUS SUBJECT?")
PAPER_TITLE = os.environ.get("PAPER_TITLE", "Where Is the Conscious Subject?")
PAPER_SUBJECT = os.environ.get("PAPER_SUBJECT", "A Dynamical Criterion for System Boundaries")
PAPER_KEYWORDS = os.environ.get(
    "PAPER_KEYWORDS",
    "consciousness; system boundaries; dynamical autonomy; causal closure; system individuation",
)
PAPER_TABLE_PROFILE = os.environ.get("PAPER_TABLE_PROFILE", "")
PAPER_PREVIEW_PAGE_BREAK = os.environ.get("PAPER_PREVIEW_PAGE_BREAK", "0") == "1"

TABLE_WIDTHS = {
    ("Field", "Required entry"): (0.25, 0.75),
    ("Status", "Formal condition", "Interpretation"): (0.20, 0.30, 0.50),
    (
        "Term or symbol",
        "Role in the framework",
        "Interpretive guardrail",
    ): (0.22, 0.28, 0.50),
    (
        "Stage",
        "Required action",
        "Failure output",
        "Guardrail",
    ): (0.12, 0.30, 0.20, 0.38),
    (
        "Case",
        "Competing boundaries",
        "Decisive perturbation",
        "Likely error if boundary is assumed",
    ): (0.15, 0.24, 0.27, 0.34),
    ("Misuse", "Why invalid", "Required correction"): (0.26, 0.35, 0.39),
    (
        "Regime and candidate",
        "",
        "",
        "Role profile",
        "Boundary result",
    ): (0.17, 0.085, 0.085, 0.28, 0.38),
    ("Candidate or test", "Autonomy/role result", "Interpretation"): (0.24, 0.32, 0.44),
    (
        "Candidate",
        "interval",
        "interval",
        "Limiting role interval",
        "Decision",
    ): (0.15, 0.16, 0.17, 0.22, 0.30),
    ("Gate", "Pass rule", "Fail rule", "Otherwise"): (0.16, 0.29, 0.28, 0.27),
    ("Distinction", "Does not establish", "Required test"): (0.22, 0.34, 0.44),
    ("Term or symbol", "Definition", "Guardrail"): (0.22, 0.34, 0.44),
    ("Failure mode", "Why spread metrics are fooled", "Corrective control"): (0.23, 0.37, 0.40),
    ("Case", "Candidate recipient classes", "Report-independent intervention", "Interpretive risk"): (0.15, 0.27, 0.30, 0.28),
    ("Recipient class", "Effect LCB", "Latency UCB", "Selectivity", "Result"): (0.25, 0.14, 0.15, 0.16, 0.30),
    ("Architecture", "Reached nodes", "Functional classes", "Diversity score", "Availability result"): (0.20, 0.14, 0.20, 0.16, 0.30),
    ("Presence result", "Character-map result", "Licensed conclusion", "Prohibited conclusion"): (0.15, 0.17, 0.35, 0.33),
    ("Character feature", "Geometric or dynamical representation", "Required guardrail"): (0.18, 0.43, 0.39),
    ("Validity question", "Pass evidence", "Failure consequence"): (0.28, 0.39, 0.33),
    ("Test item", "Phenomenal relation", "Predicted distance", "Absolute error", "Result"): (0.20, 0.15, 0.18, 0.13, 0.34),
    ("Domain", "Promising structure", "Main confound", "Required extension"): (0.16, 0.31, 0.24, 0.29),
}

if PAPER_TABLE_PROFILE == "availability":
    TABLE_WIDTHS.update(
        {
            ("Distinction", "Does not establish", "Required test"): (0.18, 0.36, 0.46),
            ("Status", "Formal condition", "Interpretation"): (0.15, 0.48, 0.37),
            ("Term or symbol", "Definition", "Guardrail"): (0.19, 0.34, 0.47),
            ("Stage", "Required action", "Failure output", "Guardrail"): (0.13, 0.29, 0.24, 0.34),
            ("Field", "Required entry"): (0.23, 0.77),
            ("Failure mode", "Why spread metrics are fooled", "Corrective control"): (0.20, 0.39, 0.41),
            ("Case", "Candidate recipient classes", "Report-independent intervention", "Interpretive risk"): (0.13, 0.29, 0.31, 0.27),
            ("Recipient class", "Effect LCB", "Latency UCB", "Selectivity", "Result"): (0.26, 0.12, 0.13, 0.12, 0.37),
            ("Architecture", "Reached nodes", "Functional classes", "Diversity score", "Availability result"): (0.21, 0.11, 0.18, 0.14, 0.36),
            ("Gate", "Pass rule", "Fail rule", "Otherwise"): (0.17, 0.31, 0.29, 0.23),
        }
    )

if PAPER_TABLE_PROFILE == "character":
    TABLE_WIDTHS.update(
        {
            ("Presence result", "Character-map result", "Licensed conclusion", "Prohibited conclusion"): (0.15, 0.18, 0.35, 0.32),
            ("Character feature", "Geometric or dynamical representation", "Required guardrail"): (0.17, 0.45, 0.38),
            ("Status", "Formal condition", "Interpretation"): (0.15, 0.47, 0.38),
            ("Term or symbol", "Definition", "Guardrail"): (0.20, 0.34, 0.46),
            ("Stage", "Required action", "Failure output", "Guardrail"): (0.15, 0.30, 0.22, 0.33),
            ("Field", "Required entry"): (0.22, 0.78),
            ("Validity question", "Pass evidence", "Failure consequence"): (0.29, 0.42, 0.29),
            ("Test item", "Phenomenal relation", "Predicted distance", "Absolute error", "Result"): (0.14, 0.16, 0.17, 0.14, 0.39),
            ("Domain", "Promising structure", "Main confound", "Required extension"): (0.15, 0.32, 0.23, 0.30),
        }
    )

if PAPER_TABLE_PROFILE == "ablation":
    TABLE_WIDTHS.update(
        {
            ("Ablation type", "Operation", "Question answered", "What it cannot establish alone"): (0.15, 0.29, 0.29, 0.27),
            ("Status", "Formal or evidential condition", "Licensed interpretation"): (0.14, 0.53, 0.33),
            ("Term or symbol", "Definition", "Guardrail"): (0.18, 0.37, 0.45),
            ("Stage", "Required action", "Circularity blocked", "Failure output"): (0.12, 0.28, 0.38, 0.22),
            ("Field", "Required entry"): (0.23, 0.77),
            ("Audit question", "Pass evidence", "Failure consequence"): (0.30, 0.46, 0.24),
            ("Case", "V", "N1", "N2", "N3", "Diagnostic status"): (0.11, 0.07, 0.07, 0.07, 0.07, 0.61),
            ("Component", "Apparent ablation", "Principal confound", "Required control"): (0.12, 0.24, 0.30, 0.34),
        }
    )


def run(*args: str) -> None:
    subprocess.run(args, check=True)


def set_font(style, name: str, size: int, bold=None, italic=None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_element = paragraph.add_run()._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run_element.extend([begin, instruction, end])


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_width = tc_pr.find(qn("w:tcW"))
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        tc_pr.append(tc_width)
    tc_width.set(qn("w:type"), "dxa")
    tc_width.set(qn("w:w"), str(int(width_inches * 1440)))
    cell.width = Inches(width_inches)


def set_cell_margins(cell, *, top: int, bottom: int, start: int = 90, end: int = 90) -> None:
    """Set cell insets in twentieths of a point (DXA)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        margin = tc_mar.find(qn(f"w:{side}"))
        if margin is None:
            margin = OxmlElement(f"w:{side}")
            tc_mar.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def add_word_math_label(paragraph, base: str, subscript: str) -> None:
    """Add a compact math-like label using ordinary Word runs."""
    base_run = paragraph.add_run(base)
    base_run.italic = True
    subscript_run = paragraph.add_run(subscript)
    subscript_run.font.subscript = True


def replace_g1_header_math(table) -> None:
    """Avoid dark OMML equations on the navy DOCX header background."""
    metric_specs = ((1, "J", "self"), (2, "A", "Θ"))
    for column_index, base, subscript in metric_specs:
        cell = table.rows[0].cells[column_index]
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_word_math_label(paragraph, base, subscript)

    profile_cell = table.rows[0].cells[3]
    profile_cell.text = ""
    paragraph = profile_cell.paragraphs[0]
    paragraph.add_run("Role profile\n(")
    for index, subscript in enumerate(("I", "A", "R", "V")):
        add_word_math_label(paragraph, "F", subscript)
        paragraph.add_run(", " if index < 3 else ")")


def replace_g3_header_math(table) -> None:
    """Render G.3 metric headers as styleable Word runs."""
    for column_index, base, subscript, suffix in (
        (1, "A", "Θ", " interval"),
        (2, "J", "self", " interval"),
    ):
        cell = table.rows[0].cells[column_index]
        cell.text = ""
        paragraph = cell.paragraphs[0]
        add_word_math_label(paragraph, base, subscript)
        paragraph.add_run(suffix)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def style_docx() -> None:
    document = Document(RAW_DOCX)

    for section in document.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)
        header = section.header.paragraphs[0]
        header.text = PAPER_HEADER
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.runs[0].font.name = "Times New Roman"
        header.runs[0].font.size = Pt(9)
        header.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        add_page_number(section.footer.paragraphs[0])

    base_styles = ["Normal", "Body Text", "First Paragraph", "Compact"]
    for name in base_styles:
        if name in document.styles:
            style = document.styles[name]
            set_font(style, "Times New Roman", 11)
            style.paragraph_format.line_spacing = 1.15
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.widow_control = True

    title = document.styles["Title"]
    set_font(title, "Times New Roman", 22, bold=True)
    title.font.color.rgb = RGBColor(28, 45, 64)
    title.paragraph_format.space_before = Pt(80)
    title.paragraph_format.space_after = Pt(8)

    subtitle = document.styles["Subtitle"]
    set_font(subtitle, "Times New Roman", 14, italic=True)
    subtitle.font.color.rgb = RGBColor(66, 77, 88)
    subtitle.paragraph_format.space_after = Pt(28)

    if "Author" in document.styles:
        author = document.styles["Author"]
        set_font(author, "Times New Roman", 12)
        author.paragraph_format.space_after = Pt(8)
    if "Date" in document.styles:
        date = document.styles["Date"]
        set_font(date, "Times New Roman", 10)
        date.paragraph_format.line_spacing = 1.15

    heading_specs = {
        "Heading 1": (15, 12, 6),
        "Heading 2": (13, 14, 6),
        "Heading 3": (11, 10, 4),
    }
    for name, (size, before, after) in heading_specs.items():
        if name in document.styles:
            style = document.styles[name]
            set_font(style, "Times New Roman", size, bold=True)
            style.font.color.rgb = RGBColor(28, 45, 64)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True

    if "Block Text" in document.styles:
        block = document.styles["Block Text"]
        set_font(block, "Times New Roman", 10, italic=True)
        block.paragraph_format.left_indent = Inches(0.35)
        block.paragraph_format.right_indent = Inches(0.35)
        block.paragraph_format.space_before = Pt(6)
        block.paragraph_format.space_after = Pt(6)

    if not any(paragraph.text == "A visual preview:" for paragraph in document.paragraphs):
        for paragraph in document.paragraphs:
            if paragraph._p.xpath(".//w:drawing"):
                paragraph.insert_paragraph_before("A visual preview:")
                break

    for shape_index, shape in enumerate(document.inline_shapes):
        if shape.type is not None:
            aspect_ratio = shape.height / shape.width
            target_width = 8.5 if shape_index == 0 else 6.5
            shape.width = Inches(target_width)
            shape.height = Inches(target_width * aspect_ratio)

    drawing_paragraphs = [
        paragraph for paragraph in document.paragraphs if paragraph._p.xpath(".//w:drawing")
    ]
    for drawing_index, paragraph in enumerate(drawing_paragraphs):
        if drawing_index == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Inches(-1.0)
            paragraph.paragraph_format.right_indent = Inches(-1.0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.left_indent = Inches(0)
            paragraph.paragraph_format.right_indent = Inches(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)

    abstract_seen = False
    references_seen = False
    for paragraph in document.paragraphs:
        if paragraph.text == "A visual preview:":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if PAPER_PREVIEW_PAGE_BREAK:
                paragraph.paragraph_format.page_break_before = True
                paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_before = Pt(14)
            paragraph.paragraph_format.space_after = Pt(8)
            for run_item in paragraph.runs:
                run_item.font.name = "Times New Roman"
                run_item.font.size = Pt(11)
                run_item.font.bold = True
                run_item.font.color.rgb = RGBColor(28, 45, 64)
                run_item._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        elif paragraph.text == "Abstract":
            paragraph.paragraph_format.page_break_before = True
            abstract_seen = True
        elif paragraph.text == "References":
            paragraph.paragraph_format.page_break_before = True
            references_seen = True
            continue
        elif paragraph.text.startswith("Appendix ") and ":" in paragraph.text:
            references_seen = False
            paragraph.paragraph_format.page_break_before = True
        if references_seen and paragraph.text and paragraph.text != "References":
            if paragraph.text.startswith("Central paper and related publications:"):
                paragraph.paragraph_format.left_indent = Inches(0)
                paragraph.paragraph_format.first_line_indent = Inches(0)
                paragraph.paragraph_format.space_before = Pt(4 if PAPER_TABLE_PROFILE == "ablation" else 10)
                paragraph.paragraph_format.space_after = Pt(4 if PAPER_TABLE_PROFILE == "ablation" else 8)
            else:
                paragraph.paragraph_format.left_indent = Inches(0.3)
                paragraph.paragraph_format.first_line_indent = Inches(-0.3)
                paragraph.paragraph_format.space_after = Pt(2 if PAPER_TABLE_PROFILE == "ablation" else 8)
            if PAPER_TABLE_PROFILE == "ablation":
                paragraph.paragraph_format.line_spacing = 1.0
                for run_item in paragraph.runs:
                    run_item.font.size = Pt(7.5)

    # Keep the compact algorithm genuinely compact so it remains a single,
    # readable unit instead of leaving its final steps alone on a new page.
    in_compact_algorithm = False
    for paragraph in document.paragraphs:
        if paragraph.text == "Appendix B.1 Compact algorithm":
            in_compact_algorithm = True
            continue
        if paragraph.text == "Appendix C: Preregistration template":
            in_compact_algorithm = False
        if in_compact_algorithm and paragraph.text:
            paragraph.paragraph_format.line_spacing = 0.9 if PAPER_TABLE_PROFILE == "ablation" else 0.95
            paragraph.paragraph_format.space_after = Pt(0 if PAPER_TABLE_PROFILE == "ablation" else 1)
            for run_item in paragraph.runs:
                run_item.font.size = Pt(8.75 if PAPER_TABLE_PROFILE == "ablation" else 9.5)

    for table in document.tables:
        table.style = "Table"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table_pr = table._tbl.tblPr
        layout = table_pr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            table_pr.append(layout)
        layout.set(qn("w:type"), "fixed")

        headers = tuple(cell.text.replace("\n", " ").strip() for cell in table.rows[0].cells)
        ratios = TABLE_WIDTHS.get(headers)
        if (
            ratios is None
            and PAPER_TABLE_PROFILE == "ablation"
            and len(headers) == 6
            and headers[0] == "Case"
            and headers[-1] == "Diagnostic status"
        ):
            # Pandoc/Word can serialize the four math-only component headers
            # differently (for example, N1, N 1, or an OMML placeholder).
            # The stable textual edge columns identify Appendix E.1 without
            # depending on those serialization details.
            ratios = (0.11, 0.07, 0.07, 0.07, 0.07, 0.61)
        if ratios is None:
            ratios = tuple(1 / len(table.columns) for _ in table.columns)
        widths = tuple(6.5 * ratio for ratio in ratios)
        if headers == ("Regime and candidate", "", "", "Role profile", "Boundary result"):
            replace_g1_header_math(table)
        elif headers == (
            "Candidate",
            "interval",
            "interval",
            "Limiting role interval",
            "Decision",
        ):
            replace_g3_header_math(table)

        for row_index, row in enumerate(table.rows):
            prevent_row_split(row)
            if row_index == 0:
                set_repeat_header(row)
            fill = "1C2D40" if row_index == 0 else ("EAF1F6" if row_index % 2 == 0 else "FFFFFF")
            for column_index, cell in enumerate(row.cells):
                set_cell_width(cell, widths[column_index])
                set_cell_shading(cell, fill)
                if row_index == 0:
                    set_cell_margins(cell, top=100, bottom=55)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(3)
                    if row_index == 0:
                        paragraph.paragraph_format.line_spacing = 0.94
                        paragraph.paragraph_format.space_after = Pt(1)
                    for run_item in paragraph.runs:
                        run_item.font.name = "Times New Roman"
                        if PAPER_TABLE_PROFILE == "ablation" and headers == ("Term or symbol", "Definition", "Guardrail"):
                            run_item.font.size = Pt(8.25)
                        else:
                            run_item.font.size = Pt(9)
                        run_item._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                        if row_index == 0:
                            run_item.bold = True
                            run_item.font.color.rgb = RGBColor(255, 255, 255)

    document.core_properties.title = PAPER_TITLE
    document.core_properties.subject = PAPER_SUBJECT
    document.core_properties.author = "Phil Stilwell"
    document.core_properties.keywords = PAPER_KEYWORDS
    DOC_OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(DOC_OUT)


def main() -> None:
    TMP.mkdir(parents=True, exist_ok=True)
    PDF_OUT.parent.mkdir(parents=True, exist_ok=True)
    preview = SOURCE.parent / "visual-preview.png"
    pdf_source = SOURCE.read_text(encoding="utf-8")
    if preview.exists():
        pdf_source = pdf_source.replace(
            "![](visual-preview.png){width=100%}",
            f"\\noindent\\makebox[\\textwidth][c]{{\\includegraphics[width=\\paperwidth]{{{preview.as_posix()}}}}}",
        )
    PDF_SOURCE.write_text(pdf_source, encoding="utf-8")
    run(
        "pandoc",
        str(SOURCE),
        "--from=markdown+tex_math_dollars+tex_math_single_backslash",
        f"--resource-path={SOURCE.parent}",
        f"--lua-filter={TABLE_FILTER}",
        "--to=docx",
        f"--output={RAW_DOCX}",
    )
    style_docx()
    run(
        "pandoc",
        str(PDF_SOURCE),
        "--from=markdown+tex_math_dollars+tex_math_single_backslash",
        f"--resource-path={SOURCE.parent}",
        f"--lua-filter={TABLE_FILTER}",
        "--pdf-engine=xelatex",
        "-V",
        "papersize=letter",
        "-V",
        "geometry:margin=0.9in",
        "-V",
        "fontsize=11pt",
        "-V",
        "mainfont=Times New Roman",
        "-V",
        "sansfont=Arial",
        "-V",
        "header-includes=\\usepackage{graphicx}\\usepackage{colortbl}\\definecolor{TableHeader}{HTML}{1C2D40}\\definecolor{TableAlt}{HTML}{EAF1F6}",
        "-V",
        "linestretch=1.15",
        "-V",
        "colorlinks=true",
        "-V",
        "linkcolor=NavyBlue",
        "-V",
        "urlcolor=NavyBlue",
        f"--output={PDF_OUT}",
    )


if __name__ == "__main__":
    main()
