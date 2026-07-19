#!/usr/bin/env python3
"""Build and validate the Cø / N* student textbook.

The source stays human-editable Markdown. Pandoc supplies semantic structure,
python-docx supplies the student-edition design system, and LibreOffice creates
the tagged PDF. The release is published only after structural and artifact QA.
"""

from __future__ import annotations

import fcntl
import hashlib
import json
import math
import os
import re
import shutil
import statistics
import subprocess
import sys
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader


PROJECT = Path(__file__).resolve().parents[1]
SOURCE_DIR = PROJECT / "student-textbook"
FRONTMATTER = SOURCE_DIR / "frontmatter.md"
CHAPTER_SOURCES = (
    SOURCE_DIR / "drafts" / "chapters-01-05.md",
    SOURCE_DIR / "drafts" / "chapters-06-10.md",
    SOURCE_DIR / "drafts" / "chapters-11-14-and-backmatter.md",
)
CORPUS_SOURCES = (
    PROJECT / "public" / "paper.pdf",
    PROJECT / "papers" / "where-is-the-conscious-subject" / "manuscript.md",
    PROJECT / "papers" / "consciousness-without-report" / "manuscript.md",
    PROJECT / "papers" / "indeterminacy-as-scientific-result" / "manuscript.md",
    PROJECT / "papers" / "ablating-n-star" / "manuscript.md",
    PROJECT / "papers" / "from-phenomenal-presence-to-phenomenal-character" / "manuscript.md",
    PROJECT / "papers" / "consciousness-in-the-schematic" / "manuscript.md",
)
ASSEMBLED_SOURCE = SOURCE_DIR / "student-textbook.md"
DOCX_OUT = PROJECT / "output" / "doc" / "learning-c0-n-star-student-textbook.docx"
PDF_OUT = PROJECT / "output" / "pdf" / "learning-c0-n-star-student-textbook.pdf"
MANIFEST_OUT = PROJECT / "output" / "student-textbook-build-manifest.json"

TITLE = "Learning Cø / N*: A Student Textbook of Phenomenal Presence"
EDITION = "1.1"
CORPUS_BASELINE = "July 18, 2026"
BODY_FONT = os.environ.get("C0_TEXTBOOK_BODY_FONT", "Liberation Serif")
DISPLAY_FONT = os.environ.get("C0_TEXTBOOK_DISPLAY_FONT", "Liberation Sans")
PUBLIC_PDF_OUT = (
    PROJECT
    / "public"
    / "teaching"
    / "student-textbook"
    / EDITION
    / PDF_OUT.name
)
ALLOW_PUBLIC_REPLACE_ENV = "C0_TEXTBOOK_ALLOW_PUBLIC_REPLACE"

# Exact samples from the user-supplied five-swatch image.
BURGUNDY = "601D1F"
SANDSTONE = "AA9062"
ESPRESSO = "3B2317"
UMBER = "8A5C39"
PALE_GOLD = "FBE4AA"

# Palette-derived paper tints used for long-form readability.
PAPER = "FFFDF8"
CREAM = "FBF3DE"
BLUSH = "F4E5E1"
MUSHROOM = "EEE7DF"
WHITE = "FFFFFF"
TEXT = RGBColor(59, 35, 23)

EXPECTED_CHAPTERS = {
    1: "The Explanandum and the Theory Landscape",
    2: "The Core Biconditional and Network-Dynamics Model",
    3: "Where Is the Conscious Subject?",
    4: "Viability and Integration: Realization and Unity",
    5: "Consciousness Without Report: System-Wide Availability",
    6: "Recurrent Stability and Temporal Presence",
    7: "What Does the Evidence License?",
    8: "Does Every Conjunct Earn Its Place?",
    9: "From Phenomenal Presence to Phenomenal Character",
    10: "Reading Consciousness from the Schematic",
    11: "Biological and Clinical Applications",
    12: "Nonhuman Animals, Organoids, and Artificial Systems",
    13: "Differentiation, Validation, and the Strongest Objections",
    14: "Capstone Adversarial Research Design",
}

CHAPTER_PATTERN = re.compile(
    r"^#\s+Chapter\s+(\d+)\s+-\s+(.+?)\s*$", re.MULTILINE
)
FORBIDDEN_PATTERNS = {
    "instructor key": re.compile(r"\binstructor\s+key\b", re.I),
    "answer guide": re.compile(r"\banswer\s+guide\b", re.I),
    "reveal sheet": re.compile(r"\breveal\s+sheet\b", re.I),
    "controlled reveal": re.compile(r"\bcontrolled[- ]reveal\b", re.I),
    "provided solution": re.compile(r"\b(?:correct|model)\s+answer\s*:", re.I),
    "placeholder": re.compile(r"\b(?:TODO|TBD|FIXME|lorem ipsum)\b", re.I),
}

# The document/PDF production policy uses ASCII hyphens only. Reject Unicode
# dash variants because they are visually confusable and layout-sensitive.
UNSAFE_UNICODE_DASHES = {
    "\u058a": "ARMENIAN HYPHEN",
    "\u1806": "MONGOLIAN TODO SOFT HYPHEN",
    "\u2010": "HYPHEN",
    "\u2011": "NON-BREAKING HYPHEN",
    "\u2012": "FIGURE DASH",
    "\u2013": "EN DASH",
    "\u2014": "EM DASH",
    "\u2015": "HORIZONTAL BAR",
    "\u2043": "HYPHEN BULLET",
    "\u2e17": "DOUBLE OBLIQUE HYPHEN",
    "\ufe58": "SMALL EM DASH",
    "\ufe63": "SMALL HYPHEN-MINUS",
    "\uff0d": "FULLWIDTH HYPHEN-MINUS",
}

SOURCE_CALLOUT_PATTERNS = {
    "Source note": re.compile(r"^\*\*Source note\.\*\*\s+\S.*$", re.I | re.M),
    "Context note": re.compile(r"^\*\*Context note\.\*\*\s+\S.*$", re.I | re.M),
    "Research-status note": re.compile(
        r"^\*\*Research-status note\.\*\*\s+\S.*$", re.I | re.M
    ),
}

THOUGHT_EXPERIMENT_HEADING = re.compile(
    r"^###\s+Thought experiment:\s+\S.+\s*$", re.I | re.M
)
PRACTICE_HEADING_PATTERNS = {
    "Quick checks": re.compile(r"^###\s+Quick checks\s*$", re.I | re.M),
    "Individual transfer": re.compile(
        r"^###\s+Individual(?:\s+transfer(?:\s+task)?)?(?::\s*.+)?\s*$",
        re.I | re.M,
    ),
    "Collaborative work": re.compile(
        r"^###\s+Collaborative\s+(?:exercise(?:\s+and\s+discussion)?|discussion)"
        r"(?::\s*.+)?\s*$",
        re.I | re.M,
    ),
    "Counterexample challenge": re.compile(
        r"^###\s+Counterexample challenge\s*$", re.I | re.M
    ),
    "Exit ticket": re.compile(r"^###\s+Exit ticket\s*$", re.I | re.M),
}

STANDALONE_WORKSHEET_HEADINGS = {
    "Blank notation translation",
    "Empirical-claim annotation sheet",
    "Reading-record template",
}
FORCED_PAGE_BREAK_HEADINGS = {
    "2. Reconstruct before criticizing",
    "4. Distinguish kinds of disagreement",
    "Discussion norms for difficult topics",
}
APPENDIX_D_SECTION = re.compile(r"^(?:Cover(?:\s+record)?|[A-R]\.\s+\S.+)$", re.I)


def run(*args: str, cwd: Path | None = None) -> subprocess.CompletedProcess[str]:
    completed = subprocess.run(
        args,
        cwd=cwd or PROJECT,
        check=False,
        capture_output=True,
        text=True,
    )
    if completed.returncode:
        raise RuntimeError(
            f"Command failed ({completed.returncode}): {' '.join(args)}\n"
            f"stdout:\n{completed.stdout}\nstderr:\n{completed.stderr}"
        )
    return completed


def find_tool(name: str, env_name: str) -> str:
    override = os.environ.get(env_name)
    if override:
        path = Path(override).expanduser().resolve()
        if not path.exists():
            raise FileNotFoundError(f"{env_name} does not exist: {path}")
        return str(path)
    found = shutil.which(name)
    if not found:
        raise FileNotFoundError(f"Required executable not found: {name}")
    return found


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def word_count(text: str) -> int:
    without_code = re.sub(r"```.*?```", " ", text, flags=re.S)
    return len(re.findall(r"\b[\wÀ-ÖØ-öø-ÿ*øØ]+(?:[-’'][\wÀ-ÖØ-öø-ÿ*]+)*\b", without_code))


def env_flag(name: str) -> bool:
    return os.environ.get(name, "").strip().casefold() in {"1", "true", "yes", "on"}


def normalize_heading(text: str) -> str:
    text = text.strip().rstrip("#").strip()
    text = text.replace("\\*", "*")
    text = re.sub(r"[*_`]", "", text)
    return re.sub(r"\s+", " ", text).strip()


def unsafe_unicode_dash_report(text: str, *, artifact: str) -> dict[str, object]:
    findings: list[dict[str, object]] = []
    for character, name in UNSAFE_UNICODE_DASHES.items():
        for match in re.finditer(re.escape(character), text):
            findings.append(
                {
                    "name": name,
                    "codepoint": f"U+{ord(character):04X}",
                    "line": text.count("\n", 0, match.start()) + 1,
                }
            )
    if findings:
        preview = ", ".join(
            f"{item['codepoint']} {item['name']} at line {item['line']}"
            for item in findings[:6]
        )
        raise ValueError(f"{artifact} contains unsafe Unicode dash characters: {preview}")
    return {
        "status": "passed",
        "rejected_codepoints": [
            f"U+{ord(character):04X}" for character in UNSAFE_UNICODE_DASHES
        ],
        "policy": "ASCII hyphens only; Unicode dash variants are rejected",
    }


def relative_luminance(hex_color: str) -> float:
    channels = [int(hex_color[index : index + 2], 16) / 255 for index in (0, 2, 4)]
    linear = [
        channel / 12.92
        if channel <= 0.04045
        else ((channel + 0.055) / 1.055) ** 2.4
        for channel in channels
    ]
    return 0.2126 * linear[0] + 0.7152 * linear[1] + 0.0722 * linear[2]


def contrast_ratio(foreground: str, background: str) -> float:
    lighter, darker = sorted(
        (relative_luminance(foreground), relative_luminance(background)), reverse=True
    )
    return (lighter + 0.05) / (darker + 0.05)


def validate_palette_contrast() -> dict[str, object]:
    checks = (
        ("body text on paper", ESPRESSO, PAPER, 7.0),
        ("title text on paper", BURGUNDY, PAPER, 7.0),
        ("heading text on pale gold", BURGUNDY, PALE_GOLD, 7.0),
        ("table-header text", PALE_GOLD, ESPRESSO, 7.0),
        ("secondary heading text", UMBER, PALE_GOLD, 4.5),
        ("body text on sandstone", ESPRESSO, SANDSTONE, 4.5),
    )
    results = []
    for label, foreground, background, minimum in checks:
        ratio = contrast_ratio(foreground, background)
        if ratio + 1e-9 < minimum:
            raise ValueError(
                f"Palette contrast failure for {label}: {ratio:.2f}:1 < {minimum:.1f}:1"
            )
        results.append(
            {
                "use": label,
                "foreground": f"#{foreground}",
                "background": f"#{background}",
                "ratio": round(ratio, 2),
                "minimum": minimum,
            }
        )
    return {"standard": "WCAG 2.x contrast-ratio calculation", "checks": results}


def worksheet_registry(markdown: str) -> list[dict[str, object]]:
    headings = list(re.finditer(r"^(#{1,6})\s+(.+?)\s*$", markdown, re.M))
    records: list[dict[str, object]] = []
    in_appendix_d = False
    for index, match in enumerate(headings):
        level = len(match.group(1))
        title = normalize_heading(match.group(2))
        if level == 1:
            in_appendix_d = title == "Appendix D - Student Capstone Workbook"

        kind: str | None = None
        if level == 2 and title in STANDALONE_WORKSHEET_HEADINGS:
            kind = "standalone worksheet"
        elif level == 3 and title.casefold().startswith("worksheet"):
            kind = "worksheet heading"
        elif level == 2 and in_appendix_d and APPENDIX_D_SECTION.fullmatch(title):
            kind = "Appendix D workbook page"
        if kind is None:
            continue

        next_start = headings[index + 1].start() if index + 1 < len(headings) else len(markdown)
        following = markdown[match.end() : next_start]
        lines = following.splitlines()
        while lines and (not lines[0].strip() or lines[0].strip() == r"\newpage"):
            lines.pop(0)
        introductory_lines: list[str] = []
        for line in lines:
            stripped = line.strip()
            if not stripped:
                break
            if re.match(r"^(?:[-+*]\s|\d+[.)]\s|>|\||```|~~~)", stripped):
                break
            introductory_lines.append(stripped)
        introduction = " ".join(introductory_lines)
        intro_words = word_count(re.sub(r"[*_`]", "", introduction))
        records.append(
            {
                "title": title,
                "level": level,
                "kind": kind,
                "source_line": markdown.count("\n", 0, match.start()) + 1,
                "introduction_words": intro_words,
            }
        )

    duplicate_titles = sorted(
        title
        for title in {record["title"] for record in records}
        if sum(record["title"] == title for record in records) > 1
    )
    if duplicate_titles:
        raise ValueError(f"Worksheet headings are not unique: {duplicate_titles}")
    weak_introductions = [
        record for record in records if int(record["introduction_words"]) < 8
    ]
    if weak_introductions:
        details = ", ".join(
            f"{record['title']} (line {record['source_line']})"
            for record in weak_introductions
        )
        raise ValueError(f"Worksheet needs an introductory explanation of at least 8 words: {details}")

    appendix_d_records = [
        record for record in records if record["kind"] == "Appendix D workbook page"
    ]
    appendix_d_labels = {
        "Cover" if str(record["title"]).casefold().startswith("cover") else str(record["title"])[0].upper()
        for record in appendix_d_records
    }
    expected_appendix_d_labels = {"Cover", *(chr(code) for code in range(ord("A"), ord("R") + 1))}
    if appendix_d_labels != expected_appendix_d_labels:
        raise ValueError(
            "Appendix D must register Cover plus Sections A-R as worksheet pages; "
            f"found labels {sorted(appendix_d_labels)}"
        )
    return records


def assemble_source() -> str:
    missing = [
        path
        for path in (FRONTMATTER, *CHAPTER_SOURCES, *CORPUS_SOURCES)
        if not path.exists()
    ]
    if missing:
        raise FileNotFoundError("Missing textbook source: " + ", ".join(map(str, missing)))
    parts = [path.read_text(encoding="utf-8").rstrip() for path in (FRONTMATTER, *CHAPTER_SOURCES)]
    # Heading styles own pagination. Removing raw page-break commands prevents
    # accidental blank leaves when authors also add a break before a chapter.
    assembled = "\n\n".join(parts) + "\n"
    assembled = re.sub(r"^\\newpage\s*$", "", assembled, flags=re.MULTILINE)
    return assembled


def chapter_sections(markdown: str) -> dict[int, tuple[str, str]]:
    matches = list(CHAPTER_PATTERN.finditer(markdown))
    sections: dict[int, tuple[str, str]] = {}
    for match in matches:
        number = int(match.group(1))
        title = match.group(2).strip().rstrip("#").strip()
        next_h1 = re.search(r"^#\s+", markdown[match.end() :], re.MULTILINE)
        end = match.end() + next_h1.start() if next_h1 else len(markdown)
        sections[number] = (title, markdown[match.start() : end])
    return sections


def validate_source(markdown: str) -> dict[str, object]:
    dash_qa = unsafe_unicode_dash_report(markdown, artifact="Assembled Markdown")
    if not re.search(
        rf'^date:\s*["\']Student edition {re.escape(EDITION)}\b', markdown, re.I | re.M
    ):
        raise ValueError(f"Front matter does not identify Student edition {EDITION}")

    sections = chapter_sections(markdown)
    if set(sections) != set(EXPECTED_CHAPTERS):
        raise ValueError(
            f"Expected chapters 1-14 exactly; found {sorted(sections)}"
        )

    chapter_qa: dict[str, object] = {}
    for number, expected_title in EXPECTED_CHAPTERS.items():
        title, body = sections[number]
        if title.casefold() != expected_title.casefold():
            raise ValueError(
                f"Chapter {number} title mismatch: expected {expected_title!r}, found {title!r}"
            )
        words = word_count(body)
        if words < 2400:
            raise ValueError(f"Chapter {number} is underdeveloped at {words} words")
        required_sections = (
            "Chapter map",
            "Practice studio",
            "Chapter summary",
            "Key terms",
        )
        for heading in required_sections:
            if not re.search(rf"^##\s+{re.escape(heading)}\s*$", body, re.I | re.M):
                raise ValueError(f"Chapter {number} is missing section: {heading}")
        if number < 14:
            if not re.search(r"^##\s+Looking ahead\s*$", body, re.I | re.M):
                raise ValueError(f"Chapter {number} is missing Looking ahead")
        elif not re.search(
            r"^##\s+(?:Looking ahead|Where inquiry goes next)\s*$", body, re.I | re.M
        ):
            raise ValueError("Chapter 14 needs a forward-looking final section")

        analogy_limits = len(re.findall(r"\*\*Where the analogy breaks\.\*\*", body, re.I))
        thought_experiments = len(THOUGHT_EXPERIMENT_HEADING.findall(body))
        if analogy_limits < 2:
            raise ValueError(f"Chapter {number} needs at least two analogy-limit annotations")
        if thought_experiments < 2:
            raise ValueError(
                f"Chapter {number} needs at least two exact Thought experiment headings"
            )

        annotations = {
            label: len(re.findall(rf"\*\*{re.escape(label)}\.\*\*", body, re.I))
            for label in ("Plain language", "Why it matters", "Do not infer", "Method note", "Checkpoint")
        }
        minimums = {
            "Plain language": 1,
            "Why it matters": 1,
            "Do not infer": 1,
            "Method note": 1,
            "Checkpoint": 1,
        }
        deficient = [label for label, minimum in minimums.items() if annotations[label] < minimum]
        if deficient:
            raise ValueError(f"Chapter {number} lacks required annotations: {', '.join(deficient)}")

        studio_match = re.search(r"^##\s+Practice studio\s*$", body, re.I | re.M)
        if studio_match is None:
            raise ValueError(f"Chapter {number} has no exact Practice studio heading")
        studio_tail = body[studio_match.end() :]
        next_h2 = re.search(r"^##\s+", studio_tail, re.M)
        studio = studio_tail[: next_h2.start()] if next_h2 else studio_tail
        practice_headings: dict[str, int] = {}
        for label, pattern in PRACTICE_HEADING_PATTERNS.items():
            count = len(pattern.findall(studio))
            if count < 1:
                raise ValueError(
                    f"Chapter {number} Practice studio needs at least one exact {label} heading; "
                    f"found {count}"
                )
            practice_headings[label] = count

        research_callouts = {
            label: len(pattern.findall(body))
            for label, pattern in SOURCE_CALLOUT_PATTERNS.items()
        }
        if sum(research_callouts.values()) < 1:
            raise ValueError(
                f"Chapter {number} needs at least one Source, Context, or Research-status note"
            )

        chapter_qa[str(number)] = {
            "title": title,
            "words": words,
            "analogy_limit_annotations": analogy_limits,
            "thought_experiment_headings": thought_experiments,
            "annotations": annotations,
            "research_callouts": research_callouts,
            "practice_headings": practice_headings,
        }

    for label, pattern in FORBIDDEN_PATTERNS.items():
        match = pattern.search(markdown)
        if match:
            line = markdown.count("\n", 0, match.start()) + 1
            raise ValueError(f"Forbidden student-edition content ({label}) at line {line}")

    required_backmatter = (
        "Appendix A - Notation at a Glance",
        "Appendix B - How to Read an Empirical Claim",
        "Appendix C - How to Argue Fairly Across Theories",
        "Appendix D - Student Capstone Workbook",
        "Appendix E - Capstone Assessment Guide",
        "Glossary",
        "References and Source Trail",
    )
    for heading in required_backmatter:
        if not re.search(rf"^#\s+{re.escape(heading)}\s*$", markdown, re.M):
            raise ValueError(f"Missing back matter: {heading}")

    corpus_coverage_tokens = (
        "Cø as N*",
        "Where Is the Conscious Subject?",
        "Consciousness Without Report",
        "Indeterminacy as a Scientific Result",
        "Ablating N*",
        "From Phenomenal Presence to Phenomenal Character",
        "Reading Consciousness from the Schematic",
    )
    missing_corpus = [token for token in corpus_coverage_tokens if token not in markdown]
    if missing_corpus:
        raise ValueError(
            "The assembled textbook does not name every paper in the source trail: "
            + ", ".join(missing_corpus)
        )

    total_words = word_count(markdown)
    if total_words < 48_000:
        raise ValueError(f"Textbook is too short for a full treatment: {total_words} words")
    research_callout_totals = {
        label: sum(
            int(chapter["research_callouts"][label])
            for chapter in chapter_qa.values()
        )
        for label in SOURCE_CALLOUT_PATTERNS
    }
    for label in ("Context note", "Research-status note"):
        if research_callout_totals[label] < 1:
            raise ValueError(f"The textbook needs at least one recognized {label}")

    worksheets = worksheet_registry(markdown)
    return {
        "total_words": total_words,
        "chapters": chapter_qa,
        "corpus_items_named": list(corpus_coverage_tokens),
        "research_callout_totals": research_callout_totals,
        "worksheets": worksheets,
        "unicode_dash_policy": dash_qa,
    }


def set_font(style, name: str, size: float, *, bold=None, italic=None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def set_run_font(run, name: str, size: float | None = None) -> None:
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), name)


def set_shading(element, fill: str) -> None:
    properties = element.get_or_add_tcPr() if element.tag == qn("w:tc") else element.get_or_add_pPr()
    shd = properties.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        properties.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, color: str, *, size: int = 10) -> None:
    """Draw a complete cell outline, including for otherwise blank writing areas."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_paragraph_left_border(paragraph, color: str, size: int = 18, space: int = 10) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)


def set_cell_margins(cell, top=70, bottom=70, start=85, end=85) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def keep_final_table_rows_together(table) -> None:
    """Prevent a table's final body row from becoming a one-row continuation."""
    if len(table.rows) < 3:
        return
    for cell in table.rows[-2].cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = True


def keep_table_together(table) -> None:
    """Keep a compact semantic table on one page when its full height permits."""
    if len(table.rows) < 2:
        return
    for row in table.rows[:-1]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True


def normalized_cell_text(cell) -> str:
    texts = cell._tc.xpath(".//w:t | .//m:t")
    return re.sub(r"\s+", " ", " ".join(node.text or "" for node in texts)).strip()


def percentile(values: list[float], fraction: float) -> float:
    if not values:
        return 0.0
    ordered = sorted(values)
    position = fraction * (len(ordered) - 1)
    lower = math.floor(position)
    upper = math.ceil(position)
    if lower == upper:
        return ordered[lower]
    return ordered[lower] + (ordered[upper] - ordered[lower]) * (position - lower)


def content_aware_column_weights(table) -> list[float]:
    columns = len(table.columns)
    scores: list[float] = []
    for column_index in range(columns):
        cell_texts = [normalized_cell_text(row.cells[column_index]) for row in table.rows]
        lengths = [len(text) for text in cell_texts if text]
        longest_words = [
            max((len(word) for word in re.findall(r"\S+", text)), default=0)
            for text in cell_texts
        ]
        header = len(cell_texts[0]) if cell_texts else 0
        score = max(
            8.0,
            header * 0.82,
            percentile(lengths, 0.78) ** 0.72 * 2.35,
            percentile(longest_words, 0.9) * 1.25,
        )
        scores.append(min(score, 62.0))

    total = sum(scores) or 1.0
    weights = [score * 100.0 / total for score in scores]
    minimum = 8.0 if columns >= 6 else 10.0 if columns >= 4 else 12.0 if columns == 3 else 20.0
    maximum = 58.0 if columns >= 3 else 76.0

    # Project onto bounded percentages while preserving the content ratios of
    # columns that have not hit a limit.
    fixed: dict[int, float] = {}
    for _ in range(columns * 2):
        changed = False
        free = [index for index in range(columns) if index not in fixed]
        remaining = 100.0 - sum(fixed.values())
        free_total = sum(scores[index] for index in free) or 1.0
        for index in list(free):
            candidate = remaining * scores[index] / free_total
            if candidate < minimum:
                fixed[index] = minimum
                changed = True
            elif candidate > maximum:
                fixed[index] = maximum
                changed = True
        if not changed:
            break
    free = [index for index in range(columns) if index not in fixed]
    remaining = 100.0 - sum(fixed.values())
    free_total = sum(scores[index] for index in free) or 1.0
    return [fixed.get(index, remaining * scores[index] / free_total) for index in range(columns)]


def semantic_column_weights(table, signature: tuple[str, ...]) -> list[float] | None:
    """Override the heuristic where formulas or writable fields need semantic room."""
    all_text = " | ".join(normalized_cell_text(cell) for row in table.rows for cell in row.cells)
    exact_profiles: dict[tuple[str, ...], list[float]] = {
        (
            "Profile",
            "Boundary and viability record",
            "Component and anchor evidence",
        ): [11, 28, 61],
        (
            "Strategy",
            "Post-intervention component record",
            "Mapping, validity, and anchor record",
        ): [18, 46, 36],
        (
            "Candidate",
            "Endogenous capacity",
            "External dependence",
            "Role completeness",
            "Stability",
            "Included in frozen family?",
        ): [14, 17, 17, 17, 11, 24],
        (
            "Construct",
            "Operational definition",
            "Estimator",
            "Calibration domain",
            "Validity envelope",
            "Main confound",
            "Stress test",
        ): [12, 20, 14, 14, 14, 14, 12],
        (
            "Reviewer concern",
            "Category",
            "Team response",
            "Change made",
            "Accepted by reviewer?",
            "Residual limitation",
        ): [16, 17, 18, 14, 16, 19],
    }
    if signature in exact_profiles:
        return exact_profiles[signature]
    if len(signature) == 2 and signature[1] in {
        "Response",
        "Entry",
        "Team entry",
        "Identity, source, and genealogy",
    }:
        return [32, 68]
    if signature == ("Scientific outcome", "Proportional practical action and rationale"):
        return [18, 82]
    if signature == ("Matched-model field", "Full model", "Primary rival"):
        return [24, 38, 38]
    if signature == ("Checkpoint", "Chair/facilitator", "Other role assignments or changes"):
        return [14, 30, 56]
    if signature == ("Collaboration dimension", "Auditable evidence", "Indicator requiring revision"):
        return [20, 43, 37]
    if signature == ("Audit item", "Yes", "No", "N/A"):
        return [82, 6, 6, 6]
    if signature == (
        "Checklist item",
        "No consequence or N/A rationale",
        "Replacement diagnostic leverage or required repair",
    ):
        return [20, 40, 40]
    if (
        len(signature) == 6
        and signature[0] == "Recipient class"
        and "S / R / C" in signature
    ):
        return [20, 7, 22, 20, 16, 15]
    if signature == ("Question", "Theory A", "Theory B"):
        return [38, 31, 31]
    if signature == ("Dependency question", "Anchor 1", "Anchor 2", "Corrective action"):
        return [34, 17, 17, 32]
    if signature == ("Symbol or term", "Student reading", "Guardrail"):
        if "Resolved" in all_text:
            return [24, 37, 39]
        if "Legible" in all_text:
            return [27, 37, 36]
        if "C0 or Cø" in all_text or "minimal phenomenal presence" in all_text:
            return [15, 34, 51]
    return None


def worksheet_minimum_body_row_height(signature: tuple[str, ...]) -> float | None:
    """Give fillable worksheet tables practical handwriting space."""
    two_column_profiles: dict[str, float] = {
        "Translation field": 0.50,
        "Science-action field": 0.85,
        "Empirical-claim field": 0.48,
        "Comparison field": 0.58,
        "Argument-map field": 0.60,
        "Diagnostic-case field": 0.55,
        "Disagreement-contract field": 0.48,
        "Rehearsal record": 0.58,
        "Reviewer record": 0.85,
        "Cover field": 0.58,
        "Project-provenance field": 0.70,
        "Collaboration field": 0.62,
        "Governance and fairness field": 0.72,
        "Branch-decision field": 0.72,
        "Target-and-scope field": 0.55,
        "Claim-license field": 0.70,
        "Candidate-record field": 0.72,
        "Boundary-decision field": 0.70,
        "Bearer-and-viability field": 0.75,
        "Diagnostic-prediction field": 0.68,
        "Diagnostic-program field": 0.62,
        "Construct-record field": 0.62,
        "Measurement-rule field": 0.45,
        "Intervention-specification field": 0.65,
        "Selectivity-and-rescue field": 0.60,
        "Causal-audit field": 0.65,
        "Control-record field": 0.68,
        "Falsification-commitment field": 0.70,
        "Anchor registry": 0.48,
        "Partition-and-blinding field": 0.68,
        "Quantitative-plan field": 0.58,
        "Reproducibility-plan field": 0.58,
        "Branch-status field": 0.70,
        "Outcome-record field": 0.72,
        "Transport-bridge field": 0.65,
        "Transport-decision field": 0.70,
        "Character-bridge foundation field": 0.62,
        "Character-bridge result field": 0.65,
        "Schematic-trial design field": 0.58,
        "Schematic-trial outcome field": 0.62,
        "Ethics-and-feasibility field": 0.60,
        "Red-team record field": 0.80,
        "Final-claim field": 1.00,
        "Reading-record field": 0.48,
    }
    if len(signature) == 2 and signature[0] in two_column_profiles:
        return two_column_profiles[signature[0]]
    profiles: dict[tuple[str, ...], float] = {
        ("Transition diagram workspace",): 4.25,
        ("Causal diagram - draw here or attach a labeled page",): 4.50,
        ("Question", "Theory A", "Theory B"): 0.62,
        ("Checkpoint", "Chair/facilitator", "Other role assignments or changes"): 0.78,
        ("Matched-model field", "Full model", "Primary rival"): 0.62,
        ("Scientific outcome", "Proportional practical action and rationale"): 1.20,
        ("Audit item", "Yes", "No", "N/A"): 0.38,
        (
            "Candidate",
            "Endogenous capacity",
            "External dependence",
            "Role completeness",
            "Stability",
            "Included in frozen family?",
        ): 0.62,
        (
            "Diagnostic case",
            "Full-model prediction",
            "Rival prediction",
            "Why they differ",
            "Observation that adjudicates",
        ): 0.68,
        (
            "Construct",
            "Operational definition",
            "Estimator",
            "Calibration domain",
            "Validity envelope",
            "Main confound",
            "Stress test",
        ): 0.55,
        (
            "Control",
            "Confound it detects",
            "Expected pattern",
            "Failure consequence",
        ): 0.58,
        ("Dependency question", "Anchor 1", "Anchor 2", "Corrective action"): 0.46,
        (
            "Observed pattern",
            "License checks",
            "Study disposition",
            "Branch-specific output(s)",
            "Claim affected",
            "Required revision or next study",
        ): 0.42,
        ("Checklist item", "No consequence or N/A rationale", "Replacement diagnostic leverage or required repair"): 0.80,
    }
    return profiles.get(signature)


def set_fixed_table_columns(table, weights: list[float], total_inches: float = 5.44) -> None:
    total_twips = int(total_inches * 1440)
    raw = [total_twips * weight / sum(weights) for weight in weights]
    widths = [int(value) for value in raw]
    remainder = total_twips - sum(widths)
    for _, index in sorted(
        ((raw[index] - widths[index], index) for index in range(len(widths))), reverse=True
    )[:remainder]:
        widths[index] += 1

    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(index, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, text, end):
        run._r.append(element)
    set_run_font(run, DISPLAY_FONT, 8.5)
    run.font.color.rgb = RGBColor(138, 92, 57)


def insert_palette_strip(document: Document) -> None:
    date_paragraph = next(
        (paragraph for paragraph in document.paragraphs if paragraph.style.name == "Date"),
        document.paragraphs[3],
    )
    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    row = table.rows[0]
    row.height = Inches(0.28)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for cell, color in zip(row.cells, (BURGUNDY, SANDSTONE, ESPRESSO, UMBER, PALE_GOLD)):
        set_shading(cell._tc, color)
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
        cell.width = Inches(1.02)
        cell.text = ""
    set_fixed_table_columns(table, [20, 20, 20, 20, 20], 5.1)
    date_paragraph._p.addnext(table._tbl)


def set_document_background(document: Document) -> None:
    background = document._element.find(qn("w:background"))
    if background is None:
        background = OxmlElement("w:background")
        document._element.insert(0, background)
    background.set(qn("w:color"), PAPER)


def style_document(raw_docx: Path, destination: Path) -> dict[str, object]:
    document = Document(raw_docx)
    document.core_properties.title = TITLE
    document.core_properties.subject = (
        "Accessible graduate student textbook for the Cø / N* theory of phenomenal presence"
    )
    document.core_properties.author = "Phil Stilwell"
    document.core_properties.last_modified_by = "Cø / N* student-textbook build"
    document.core_properties.comments = (
        f"Student edition {EDITION}; corpus baseline {CORPUS_BASELINE}; no instructor keys"
    )
    document.core_properties.keywords = (
        "consciousness; phenomenal presence; N*; student textbook; integration; "
        "availability; recurrence; thought experiments; collaborative learning"
    )
    set_document_background(document)

    settings = document.settings._element
    theme = settings.find(qn("w:themeFontLang"))
    if theme is None:
        theme = OxmlElement("w:themeFontLang")
        settings.append(theme)
    theme.set(qn("w:val"), "en-US")

    for section in document.sections:
        section.page_width = Inches(7)
        section.page_height = Inches(10)
        section.top_margin = Inches(0.70)
        section.bottom_margin = Inches(0.70)
        section.left_margin = Inches(0.74)
        section.right_margin = Inches(0.74)
        section.gutter = Inches(0.08)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.32)
        section.different_first_page_header_footer = True

        header = section.header.paragraphs[0]
        header.text = "LEARNING Cø / N*  ·  STUDENT EDITION"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.paragraph_format.space_after = Pt(0)
        for run in header.runs:
            set_run_font(run, DISPLAY_FONT, 7.8)
            run.font.bold = True
            run.font.color.rgb = RGBColor(138, 92, 57)

        footer = section.footer.paragraphs[0]
        add_page_number(footer)
        section.first_page_header.paragraphs[0].text = ""
        section.first_page_footer.paragraphs[0].text = ""

    for name in ("Normal", "Body Text", "First Paragraph", "Compact"):
        if name in document.styles:
            style = document.styles[name]
            set_font(style, BODY_FONT, 10.1)
            style.font.color.rgb = TEXT
            style.paragraph_format.line_spacing = 1.035
            style.paragraph_format.space_after = Pt(3.8)
            style.paragraph_format.widow_control = True

    title = document.styles["Title"]
    set_font(title, DISPLAY_FONT, 29, bold=True)
    title.font.color.rgb = RGBColor(96, 29, 31)
    title.paragraph_format.space_before = Pt(108)
    title.paragraph_format.space_after = Pt(9)

    if "Subtitle" in document.styles:
        subtitle = document.styles["Subtitle"]
        set_font(subtitle, BODY_FONT, 15, italic=True)
        subtitle.font.color.rgb = RGBColor(138, 92, 57)
        subtitle.paragraph_format.space_after = Pt(30)
    if "Author" in document.styles:
        author = document.styles["Author"]
        set_font(author, DISPLAY_FONT, 10.5, bold=True)
        author.font.color.rgb = RGBColor(59, 35, 23)
        author.paragraph_format.space_after = Pt(4)
    if "Date" in document.styles:
        date = document.styles["Date"]
        set_font(date, DISPLAY_FONT, 9)
        date.font.color.rgb = RGBColor(138, 92, 57)
        date.paragraph_format.space_after = Pt(18)

    heading_specs = {
        "Heading 1": (DISPLAY_FONT, 21, 0, 12, BURGUNDY),
        "Heading 2": (DISPLAY_FONT, 14.5, 15, 6, BURGUNDY),
        "Heading 3": (DISPLAY_FONT, 11.4, 10, 4, UMBER),
        "Heading 4": (DISPLAY_FONT, 10.2, 8, 3, ESPRESSO),
    }
    for name, (font, size, before, after, color) in heading_specs.items():
        if name not in document.styles:
            continue
        style = document.styles[name]
        set_font(style, font, size, bold=True)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        if name == "Heading 1":
            style.paragraph_format.page_break_before = True

    if "Block Text" in document.styles:
        block = document.styles["Block Text"]
        set_font(block, BODY_FONT, 10.5, italic=True)
        block.font.color.rgb = RGBColor(96, 29, 31)
        block.paragraph_format.left_indent = Inches(0.24)
        block.paragraph_format.right_indent = Inches(0.18)
        block.paragraph_format.space_before = Pt(8)
        block.paragraph_format.space_after = Pt(8)

    callout_specs = {
        "Plain language.": (PALE_GOLD, BURGUNDY),
        "Why it matters.": (CREAM, UMBER),
        "Do not infer.": (BLUSH, BURGUNDY),
        "Method note.": (MUSHROOM, ESPRESSO),
        "Materials.": (MUSHROOM, ESPRESSO),
        "Checkpoint.": (PALE_GOLD, UMBER),
        "Where the analogy breaks.": (MUSHROOM, SANDSTONE),
        "Source note.": (CREAM, SANDSTONE),
        "Context note.": (PALE_GOLD, UMBER),
        "Research-status note.": (BLUSH, BURGUNDY),
    }
    callout_counts = {label: 0 for label in callout_specs}
    worksheet_headings: list[str] = []
    chapter_count = 0
    in_appendix_d = False
    in_contextual_readings = False
    for paragraph in document.paragraphs:
        next_sibling = paragraph._p.getnext()
        if next_sibling is not None and next_sibling.tag == qn("w:tbl"):
            paragraph.paragraph_format.keep_with_next = True
        text = paragraph.text.strip()
        if paragraph.style.name == "Heading 1":
            in_appendix_d = normalize_heading(text) == "Appendix D - Student Capstone Workbook"
            set_shading(paragraph._p, PALE_GOLD if text.startswith("Chapter ") else CREAM)
            set_paragraph_left_border(paragraph, BURGUNDY, size=26, space=12)
            paragraph.paragraph_format.left_indent = Inches(0.14)
            paragraph.paragraph_format.right_indent = Inches(0.08)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(14)
            if text.startswith("Chapter "):
                chapter_count += 1
        normalized_text = normalize_heading(text)
        if normalized_text in FORCED_PAGE_BREAK_HEADINGS:
            paragraph.paragraph_format.page_break_before = True
        if paragraph.style.name == "Heading 2":
            in_contextual_readings = normalized_text == "Selected contextual readings"
        is_worksheet = (
            paragraph.style.name == "Heading 2"
            and normalized_text in STANDALONE_WORKSHEET_HEADINGS
        ) or (
            paragraph.style.name == "Heading 3"
            and normalized_text.casefold().startswith("worksheet")
        ) or (
            paragraph.style.name == "Heading 2"
            and in_appendix_d
            and APPENDIX_D_SECTION.fullmatch(normalized_text) is not None
        )
        if is_worksheet:
            paragraph.paragraph_format.page_break_before = True
            worksheet_headings.append(normalized_text)
        if paragraph.style.name == "Block Text":
            set_shading(paragraph._p, CREAM)
            set_paragraph_left_border(paragraph, BURGUNDY)
        for label, (fill, border) in callout_specs.items():
            if text.startswith(label):
                callout_counts[label] += 1
                set_shading(paragraph._p, fill)
                set_paragraph_left_border(paragraph, border)
                paragraph.paragraph_format.left_indent = Inches(0.16)
                paragraph.paragraph_format.right_indent = Inches(0.12)
                paragraph.paragraph_format.space_before = Pt(5)
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.paragraph_format.keep_together = True
                for run in paragraph.runs:
                    set_run_font(run, BODY_FONT, 9.9)
                break
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if in_contextual_readings and not paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.space_after = Pt(1.5)
            paragraph.paragraph_format.line_spacing = 0.98
            paragraph.paragraph_format.keep_together = True
            for run in paragraph.runs:
                set_run_font(run, BODY_FONT, 9.2)
        for run in paragraph.runs:
            if run.font.name is None:
                set_run_font(run, BODY_FONT)

    # Keep the final pair of items together so a lone objective or key term is
    # not stranded, while still allowing a substantial list to paginate well.
    paragraphs = document.paragraphs
    for paragraph_index, paragraph in enumerate(paragraphs):
        if not paragraph.text.strip().startswith("Where the analogy breaks."):
            continue
        for previous_index in range(paragraph_index - 1, -1, -1):
            previous = paragraphs[previous_index]
            if previous.text.strip():
                previous.paragraph_format.keep_with_next = True
                break

    index = 0
    while index < len(paragraphs):
        if not paragraphs[index]._p.xpath(".//w:numPr"):
            index += 1
            continue
        end = index
        while end < len(paragraphs) and paragraphs[end]._p.xpath(".//w:numPr"):
            paragraphs[end].paragraph_format.keep_together = True
            end += 1
        group = paragraphs[index:end]
        is_capstone_branch_selector = any(
            "Case-level presence classification" in item.text for item in group
        ) and any("Schematic-competence audit" in item.text for item in group)
        if is_capstone_branch_selector and 1 < len(group) <= 8:
            for item in group[:-1]:
                item.paragraph_format.keep_with_next = True
        elif 1 < len(group) <= 8 and sum(len(item.text) for item in group) <= 1000:
            for item in group[:-1]:
                item.paragraph_format.keep_with_next = True
        index = end

    # Consecutive thought-experiment stages are one reasoning sequence. Keep a
    # compact sequence together, or at minimum prevent a terminal stage from
    # appearing alone on the following page.
    stage_pattern = re.compile(r"^Stage\s+\d+\.", re.I)
    index = 0
    while index < len(paragraphs):
        if not stage_pattern.match(paragraphs[index].text.strip()):
            index += 1
            continue
        end = index
        while end < len(paragraphs) and stage_pattern.match(paragraphs[end].text.strip()):
            paragraphs[end].paragraph_format.keep_together = True
            end += 1
        group = paragraphs[index:end]
        if 1 < len(group) <= 6 and sum(len(item.text) for item in group) <= 1600:
            for item in group[:-1]:
                item.paragraph_format.keep_with_next = True
        elif len(group) > 1:
            group[-2].paragraph_format.keep_with_next = True
        index = end

    for paragraph_index, paragraph in enumerate(paragraphs):
        if not paragraph.text.strip().startswith("The thought experiment highlights"):
            continue
        for previous_index in range(paragraph_index - 1, -1, -1):
            previous = paragraphs[previous_index]
            if previous.text.strip():
                previous.paragraph_format.keep_with_next = True
                break

    # A short setup line belongs with the first substantive item it introduces.
    # This blocks page bottoms such as a section heading plus one lead sentence,
    # or a colon-ended prompt whose formula/list begins on the next page.
    for index, paragraph in enumerate(paragraphs[:-1]):
        text = paragraph.text.strip()
        if not text or paragraph._p.xpath(".//w:numPr"):
            continue
        next_paragraph = paragraphs[index + 1]
        next_style = next_paragraph.style.name
        if next_style == "Heading 1":
            continue
        previous = paragraphs[index - 1] if index else None
        follows_section_heading = (
            previous is not None
            and previous.style.name in {"Heading 2", "Heading 3", "Heading 4"}
            and previous.text.strip() not in {"Looking ahead", "Where inquiry goes next"}
            and len(text) <= 360
        )
        introduces_material = len(text) <= 220 and text.endswith(":")
        if follows_section_heading or introduces_material:
            paragraph.paragraph_format.keep_with_next = True

    # Key-term definitions are compact reference units; splitting one across
    # pages makes a chapter review harder to scan and can create one-line widows.
    in_key_terms = False
    current_chapter_number: int | None = None
    compact_end_chapters = {8, 10, 13}
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name == "Heading 1":
            chapter_match = re.match(r"Chapter\s+(\d+)\b", text)
            current_chapter_number = int(chapter_match.group(1)) if chapter_match else None
        if paragraph.style.name == "Heading 2":
            in_key_terms = text == "Key terms"
            if current_chapter_number in compact_end_chapters and text in {
                "Key terms",
                "Looking ahead",
            }:
                paragraph.paragraph_format.space_before = Pt(8)
                paragraph.paragraph_format.space_after = Pt(3)
            continue
        if paragraph.style.name == "Heading 1":
            in_key_terms = False
        if in_key_terms and text:
            paragraph.paragraph_format.keep_together = True
            if current_chapter_number in compact_end_chapters:
                paragraph.paragraph_format.space_after = Pt(1.5)
                paragraph.paragraph_format.line_spacing = 0.98
                for run in paragraph.runs:
                    set_run_font(run, BODY_FONT, 9.4)

    in_chapter_summary = False
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name == "Heading 2":
            in_chapter_summary = normalize_heading(text) == "Chapter summary"
            continue
        if paragraph.style.name.startswith("Heading"):
            in_chapter_summary = False
        if in_chapter_summary and text:
            paragraph.paragraph_format.keep_together = True

    # LibreOffice does not consistently preserve inherited widow control.
    # Keep ordinary short paragraphs intact; longer exposition may still flow.
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if (
            paragraph.style.name in {"Normal", "Body Text", "First Paragraph", "Compact"}
            and text
            and len(text) <= 650
        ):
            paragraph.paragraph_format.keep_together = True

    insert_palette_strip(document)

    table_profiles: list[dict[str, object]] = []
    for table_index, table in enumerate(document.tables):
        if not table.rows or not table.columns:
            continue
        # The cover palette strip has no text and has already been fixed.
        signature = tuple(normalized_cell_text(cell) for cell in table.rows[0].cells)
        if not any(signature):
            continue
        weights = semantic_column_weights(table, signature) or content_aware_column_weights(table)
        minimum_body_row_height = worksheet_minimum_body_row_height(signature)
        is_causal_drawing_box = len(signature) == 1 and signature[0] in {
            "Transition diagram workspace",
            "Causal diagram - draw here or attach a labeled page",
        }
        set_fixed_table_columns(table, weights)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        repeat_table_header(table.rows[0])
        for row_index, row in enumerate(table.rows):
            prevent_row_split(row)
            if row_index > 0 and minimum_body_row_height is not None:
                row.height = Inches(minimum_body_row_height)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell, top=60, bottom=60, start=70, end=70)
                body_fill = CREAM if is_causal_drawing_box else (
                    CREAM if row_index % 2 == 0 else PAPER
                )
                set_shading(cell._tc, ESPRESSO if row_index == 0 else body_fill)
                if is_causal_drawing_box:
                    set_cell_border(cell, SANDSTONE, size=12)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        set_run_font(run, DISPLAY_FONT, 8.4 if len(table.columns) >= 5 else 8.8)
                        run.font.color.rgb = (
                            RGBColor(251, 228, 170) if row_index == 0 else TEXT
                        )
                        if row_index == 0:
                            run.font.bold = True
        keep_final_table_rows_together(table)
        if signature == (
            "Strategy",
            "Post-intervention component record",
            "Mapping, validity, and anchor record",
        ):
            keep_table_together(table)
        table_profile = {
            "index": table_index,
            "signature": list(signature),
            "column_width_percentages": [round(value, 2) for value in weights],
        }
        if minimum_body_row_height is not None:
            table_profile["minimum_body_row_height_inches"] = minimum_body_row_height
        table_profiles.append(table_profile)

    for section in document.sections[1:]:
        if section.start_type == WD_SECTION.NEW_PAGE:
            section.different_first_page_header_footer = False

    document.save(destination)
    return {
        "chapters_styled": chapter_count,
        "callout_paragraphs": callout_counts,
        "worksheet_headings": worksheet_headings,
        "tables": table_profiles,
    }


def pandoc_docx(pandoc: str, markdown_path: Path, destination: Path) -> None:
    run(
        pandoc,
        str(markdown_path),
        "--from=markdown+tex_math_dollars+raw_tex+smart",
        "--to=docx",
        "--standalone",
        "--resource-path",
        str(PROJECT),
        "--metadata",
        "lang=en-US",
        "--output",
        str(destination),
    )


def convert_pdf(soffice: str, docx_path: Path, destination_dir: Path) -> Path:
    profile_dir = Path(tempfile.mkdtemp(prefix="c0-student-textbook-lo-"))
    try:
        result = run(
            soffice,
            f"-env:UserInstallation={profile_dir.as_uri()}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(destination_dir),
            str(docx_path),
        )
    finally:
        shutil.rmtree(profile_dir, ignore_errors=True)
    generated = destination_dir / f"{docx_path.stem}.pdf"
    if not generated.exists():
        raise FileNotFoundError(
            f"LibreOffice did not create {generated}; stdout={result.stdout!r}"
        )
    return generated


def flatten_outline(items) -> list[str]:
    titles: list[str] = []
    for item in items:
        if isinstance(item, list):
            titles.extend(flatten_outline(item))
        else:
            title = getattr(item, "title", None)
            if title:
                titles.append(str(title))
    return titles


def pdf_font_report(pdf_path: Path, pdffonts: str | None) -> dict[str, object]:
    if not pdffonts:
        return {"available": False}
    output = run(pdffonts, str(pdf_path)).stdout
    lines = [line for line in output.splitlines() if line.strip()]
    records = []
    for line in lines[2:]:
        columns = re.split(r"\s+", line.strip())
        if len(columns) < 8:
            continue
        records.append(
            {
                "name": columns[0],
                "type": columns[1],
                "embedded": columns[3].lower() == "yes",
                "subset": columns[4].lower() == "yes",
            }
        )
    if not records:
        raise ValueError("pdffonts returned no font records")
    unembedded = [record["name"] for record in records if not record["embedded"]]
    if unembedded:
        raise ValueError(f"PDF contains unembedded fonts: {unembedded}")
    return {"available": True, "all_embedded": True, "fonts": records}


def docx_text_statistics(docx_path: Path) -> dict[str, int]:
    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    math_namespace = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    with zipfile.ZipFile(docx_path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
    text_tags = {f"{{{word_namespace}}}t", f"{{{math_namespace}}}t"}
    paragraphs: list[str] = []
    for paragraph in root.iter(f"{{{word_namespace}}}p"):
        text = "".join(node.text or "" for node in paragraph.iter() if node.tag in text_tags)
        paragraphs.append(re.sub(r"\s+", " ", text).strip())
    normalized = "\n".join(paragraphs)
    with_spaces = " ".join(part for part in paragraphs if part)
    return {
        "pages": 0,
        "words": word_count(normalized),
        "paragraphs": len(paragraphs),
        "characters": len(re.sub(r"\s+", "", normalized)),
        "characters_with_spaces": len(with_spaces),
        "lines": max(1, math.ceil(len(with_spaces) / 80)),
        "tables": sum(1 for _ in root.iter(f"{{{word_namespace}}}tbl")),
    }


def update_docx_extended_properties(docx_path: Path, *, pages: int) -> dict[str, int]:
    statistics = docx_text_statistics(docx_path)
    statistics["pages"] = pages
    property_names = {
        "Pages": statistics["pages"],
        "Words": statistics["words"],
        "Paragraphs": statistics["paragraphs"],
        "Characters": statistics["characters"],
        "CharactersWithSpaces": statistics["characters_with_spaces"],
        "Lines": statistics["lines"],
    }
    namespace = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
    ET.register_namespace("", namespace)
    temporary = docx_path.with_name(f".{docx_path.name}.{os.getpid()}.metadata")
    try:
        with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(
            temporary, "w"
        ) as destination:
            for info in source.infolist():
                data = source.read(info.filename)
                if info.filename == "docProps/app.xml":
                    root = ET.fromstring(data)
                    for name, value in property_names.items():
                        node = root.find(f"{{{namespace}}}{name}")
                        if node is None:
                            node = ET.SubElement(root, f"{{{namespace}}}{name}")
                        node.text = str(value)
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                destination.writestr(info, data)
        os.replace(temporary, docx_path)
    finally:
        temporary.unlink(missing_ok=True)
    return statistics


def read_docx_extended_properties(docx_path: Path) -> dict[str, int]:
    namespace = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
    with zipfile.ZipFile(docx_path) as archive:
        root = ET.fromstring(archive.read("docProps/app.xml"))
    result: dict[str, int] = {}
    for key, name in (
        ("pages", "Pages"),
        ("words", "Words"),
        ("paragraphs", "Paragraphs"),
        ("characters", "Characters"),
        ("characters_with_spaces", "CharactersWithSpaces"),
        ("lines", "Lines"),
    ):
        node = root.find(f"{{{namespace}}}{name}")
        if node is None or not (node.text or "").isdigit():
            raise ValueError(f"DOCX extended properties lack a numeric {name} value")
        result[key] = int(node.text or "0")
    return result


def validate_docx(
    docx_path: Path,
    *,
    expected_pages: int,
    expected_worksheet_headings: Iterable[str],
) -> dict[str, object]:
    if not zipfile.is_zipfile(docx_path):
        raise ValueError("DOCX output is not a valid OPC/ZIP container")
    with zipfile.ZipFile(docx_path) as archive:
        bad_member = archive.testzip()
        if bad_member:
            raise ValueError(f"DOCX contains a corrupt member: {bad_member}")
    document = Document(docx_path)
    core = document.core_properties
    expected_core = {
        "title": TITLE,
        "author": "Phil Stilwell",
        "subject": "Accessible graduate student textbook for the Cø / N* theory of phenomenal presence",
        "language": "en-US",
    }
    found_core = {
        "title": core.title,
        "author": core.author,
        "subject": core.subject,
        "language": core.language,
    }
    if found_core != expected_core:
        raise ValueError(f"DOCX core metadata mismatch: {found_core}")
    if f"Student edition {EDITION}" not in (core.comments or ""):
        raise ValueError("DOCX description does not contain the current edition")

    chapter_headings = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.style.name == "Heading 1" and paragraph.text.strip().startswith("Chapter ")
    ]
    if len(chapter_headings) != 14:
        raise ValueError(f"DOCX has {len(chapter_headings)} chapter headings, expected 14")
    fixed_tables = 0
    for table in document.tables:
        if not table.rows:
            continue
        layout = table._tbl.tblPr.find(qn("w:tblLayout"))
        if layout is None or layout.get(qn("w:type")) != "fixed":
            raise ValueError("DOCX contains a non-fixed table layout")
        if len(table._tbl.tblGrid) != len(table.columns):
            raise ValueError("DOCX table grid does not match its column count")
        fixed_tables += 1

    expected_worksheets = list(expected_worksheet_headings)
    worksheet_breaks: dict[str, bool] = {}
    worksheet_occurrences: dict[str, int] = {}
    for paragraph in document.paragraphs:
        title = normalize_heading(paragraph.text)
        if title not in expected_worksheets:
            continue
        worksheet_occurrences[title] = worksheet_occurrences.get(title, 0) + 1
        worksheet_breaks[title] = paragraph.paragraph_format.page_break_before is True
    worksheet_failures = [
        title
        for title in expected_worksheets
        if worksheet_occurrences.get(title) != 1 or not worksheet_breaks.get(title, False)
    ]
    if worksheet_failures:
        raise ValueError(
            "DOCX worksheet headings must occur once with a direct page break: "
            + ", ".join(worksheet_failures)
        )

    computed_statistics = docx_text_statistics(docx_path)
    computed_statistics["pages"] = expected_pages
    stored_statistics = read_docx_extended_properties(docx_path)
    comparable = {
        key: value
        for key, value in computed_statistics.items()
        if key != "tables"
    }
    if stored_statistics != comparable:
        raise ValueError(
            f"DOCX extended statistics are stale: stored={stored_statistics}, "
            f"computed={comparable}"
        )
    return {
        "chapter_headings": chapter_headings,
        "paragraphs": len(document.paragraphs),
        "tables_fixed": fixed_tables,
        "sections": len(document.sections),
        "metadata": found_core,
        "extended_statistics": stored_statistics,
        "worksheet_page_breaks": worksheet_breaks,
    }


def outline_destinations(reader: PdfReader, items=None) -> list[tuple[str, int]]:
    results: list[tuple[str, int]] = []
    for item in reader.outline if items is None else items:
        if isinstance(item, list):
            results.extend(outline_destinations(reader, item))
            continue
        title = getattr(item, "title", None)
        if title is None:
            continue
        try:
            page_number = reader.get_destination_page_number(item) + 1
        except Exception as exc:
            raise ValueError(f"Could not resolve PDF outline destination {title!r}") from exc
        results.append((normalize_heading(str(title)), page_number))
    return results


def validate_pdf_worksheet_starts(
    reader: PdfReader, page_text: list[str], expected_headings: Iterable[str]
) -> dict[str, object]:
    destinations = outline_destinations(reader)
    by_title: dict[str, list[int]] = {}
    for title, page_number in destinations:
        by_title.setdefault(title, []).append(page_number)
    pages: dict[str, int] = {}
    for expected in expected_headings:
        matches = by_title.get(expected, [])
        if len(matches) != 1:
            raise ValueError(
                f"Worksheet outline entry {expected!r} resolves to pages {matches}; expected one"
            )
        page_number = matches[0]
        leading_text = re.sub(r"\s+", " ", page_text[page_number - 1]).strip()[:650]
        if expected not in normalize_heading(leading_text):
            raise ValueError(
                f"Worksheet {expected!r} is not near the start of PDF page {page_number}"
            )
        pages[expected] = page_number
    if len(set(pages.values())) != len(pages):
        raise ValueError(f"Worksheet headings do not begin on distinct PDF pages: {pages}")
    return {
        "method": "outline destination plus extracted-text leading-position check",
        "distinct_start_pages": True,
        "pages": pages,
        "limitation": (
            "This verifies a separate physical start page, not that every worksheet's "
            "writing space fits on exactly one page; rendered-page review remains required."
        ),
    }


def pdf_figure_accessibility_report(reader: PdfReader) -> dict[str, object]:
    root = reader.trailer["/Root"]
    struct_root = root.get("/StructTreeRoot")
    if struct_root is None:
        return {"figures": 0, "missing_alternative_text": 0, "pages": []}
    page_ids = {
        getattr(page.indirect_reference, "idnum", None): index + 1
        for index, page in enumerate(reader.pages)
    }
    total = 0
    missing_pages: list[int | None] = []

    def visit(value, inherited_page=None) -> None:
        nonlocal total
        if value is None or isinstance(value, (int, float, str)):
            return
        if isinstance(value, list):
            for child in value:
                visit(child, inherited_page)
            return
        try:
            node = value.get_object()
        except AttributeError:
            node = value
        if not hasattr(node, "get"):
            return
        page_reference = node.get("/Pg") or inherited_page
        if str(node.get("/S") or "") == "/Figure":
            total += 1
            if not node.get("/Alt") and not node.get("/ActualText"):
                page_id = getattr(page_reference, "idnum", None)
                missing_pages.append(page_ids.get(page_id))
        visit(node.get("/K"), page_reference)

    visit(struct_root.get_object().get("/K"))
    known_pages = sorted({page for page in missing_pages if page is not None})
    report = {
        "figures": total,
        "missing_alternative_text": len(missing_pages),
        "pages": known_pages,
        "release_gate": "hard failure",
        "limitation": (
            "The structural check verifies Alt or ActualText on Figure tags; rendered-page "
            "review must still confirm that decorative rules were not mis-tagged as figures."
        ),
    }
    if missing_pages:
        raise ValueError(
            f"PDF contains {len(missing_pages)} Figure tags without alternative text "
            f"(pages {known_pages})"
        )
    return report


def validate_pdf(
    pdf_path: Path,
    pdfinfo: str | None,
    pdffonts: str | None,
    *,
    expected_worksheet_headings: Iterable[str],
) -> dict[str, object]:
    reader = PdfReader(str(pdf_path))
    if len(reader.pages) < 100:
        raise ValueError(f"PDF is unexpectedly short at {len(reader.pages)} pages")
    page_text = [(page.extract_text() or "").strip() for page in reader.pages]
    worksheet_pages = validate_pdf_worksheet_starts(
        reader, page_text, expected_worksheet_headings
    )
    worksheet_start_pages = set(worksheet_pages["pages"].values())
    outline_start_pages = {page for _, page in outline_destinations(reader)}
    structural_context_pages = outline_start_pages | {
        page - 1 for page in outline_start_pages if page > 1
    }
    empty_pages = [index + 1 for index, text in enumerate(page_text) if not text]
    intentionally_open_pages = [
        index + 1
        for index, text in enumerate(page_text)
        if index > 0
        and word_count(text) < 100
        and (
            text.count("_") >= 10
            or text.count("☐") >= 8
            or text.count("□") >= 8
            or "Final claim the team is willing to risk" in text
            or index + 1 in worksheet_start_pages
            or index + 1 in structural_context_pages
        )
    ]
    sparse_pages = [
        index + 1
        for index, text in enumerate(page_text)
        if index > 0
        and 0 < word_count(text) < 100
        and index + 1 not in intentionally_open_pages
    ]
    if empty_pages:
        raise ValueError(f"PDF contains empty pages: {empty_pages}")
    if sparse_pages:
        sparse_previews = {
            page: re.sub(r"\s+", " ", page_text[page - 1])[:500]
            for page in sparse_pages
        }
        raise ValueError(
            f"PDF contains materially sparse pages: {sparse_pages}; "
            f"text previews: {sparse_previews}"
        )

    full_text = "\n".join(page_text)
    dash_qa = unsafe_unicode_dash_report(full_text, artifact="Extracted PDF text")
    raw_markup_patterns = {
        "raw page break": re.compile(r"\\(?:newpage|pagebreak)\b"),
        "raw TeX environment": re.compile(r"\\(?:begin|end)\{"),
        "raw TeX command": re.compile(r"\\(?:frac|operatorname)\{"),
        "Markdown heading marker": re.compile(r"^#{1,6}\s", re.M),
        "Markdown code marker": re.compile(r"`"),
        "unrendered math delimiter": re.compile(r"\$"),
    }
    leaked_markup = [label for label, pattern in raw_markup_patterns.items() if pattern.search(full_text)]
    if leaked_markup:
        raise ValueError(f"PDF contains raw source markup: {leaked_markup}")
    normalized_full_text = re.sub(r"\s+", " ", full_text)
    for number, title in EXPECTED_CHAPTERS.items():
        if f"Chapter {number}" not in normalized_full_text or title not in normalized_full_text:
            raise ValueError(f"PDF text is missing Chapter {number}: {title}")

    outline = flatten_outline(reader.outline)
    missing_outline = [
        f"Chapter {number} - {title}"
        for number, title in EXPECTED_CHAPTERS.items()
        if not any(item.startswith(f"Chapter {number}") for item in outline)
    ]
    if missing_outline:
        raise ValueError(f"PDF outline is missing chapter entries: {missing_outline}")

    root = reader.trailer["/Root"]
    mark_info = root.get("/MarkInfo")
    marked = bool(mark_info and mark_info.get_object().get("/Marked", False))
    language = str(root.get("/Lang") or "")
    if not marked or "/StructTreeRoot" not in root:
        raise ValueError("PDF is not tagged")
    if language != "en-US":
        raise ValueError(f"PDF document language is {language!r}, expected 'en-US'")

    metadata = reader.metadata
    metadata_report = {
        "title": metadata.title if metadata else None,
        "author": metadata.author if metadata else None,
        "subject": metadata.subject if metadata else None,
    }
    if metadata_report != {
        "title": TITLE,
        "author": "Phil Stilwell",
        "subject": "Accessible graduate student textbook for the Cø / N* theory of phenomenal presence",
    }:
        raise ValueError(f"PDF metadata mismatch: {metadata_report}")

    expected_box = (0.0, 0.0, 504.0, 720.0)
    geometry_failures: list[int] = []
    rotation_failures: list[int] = []
    for index, page in enumerate(reader.pages, start=1):
        media_box = tuple(float(value) for value in page.mediabox)
        crop_box = tuple(float(value) for value in page.cropbox)
        if any(abs(found - expected) > 0.1 for found, expected in zip(media_box, expected_box)):
            geometry_failures.append(index)
        elif any(abs(found - expected) > 0.1 for found, expected in zip(crop_box, expected_box)):
            geometry_failures.append(index)
        if (page.rotation or 0) % 360 != 0:
            rotation_failures.append(index)
    if geometry_failures:
        raise ValueError(f"PDF pages have unexpected geometry: {geometry_failures}")
    if rotation_failures:
        raise ValueError(f"PDF pages have unexpected rotation: {rotation_failures}")

    info_report: dict[str, str] = {}
    if pdfinfo:
        for line in run(pdfinfo, str(pdf_path)).stdout.splitlines():
            if ":" in line:
                key, value = line.split(":", 1)
                info_report[key.strip()] = value.strip()

    return {
        "pages": len(reader.pages),
        "empty_pages": empty_pages,
        "sparse_pages": sparse_pages,
        "intentionally_open_pages": intentionally_open_pages,
        "outline_entries": len(outline),
        "tagged": marked,
        "language": language,
        "metadata": metadata_report,
        "page_geometry_points": list(expected_box),
        "rotation_degrees": 0,
        "raw_markup": {"status": "passed", "patterns_checked": list(raw_markup_patterns)},
        "unicode_dash_policy": dash_qa,
        "worksheet_pages": worksheet_pages,
        "figure_accessibility": pdf_figure_accessibility_report(reader),
        "pdfinfo": info_report,
        "fonts": pdf_font_report(pdf_path, pdffonts),
    }


def tool_version(executable: str, *args: str) -> str:
    completed = subprocess.run(
        (executable, *args), check=False, capture_output=True, text=True
    )
    output = (completed.stdout or completed.stderr).strip().splitlines()
    return output[0] if output else "unknown"


def publish(staged: Path, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    temporary = destination.with_name(f".{destination.name}.{os.getpid()}.publishing")
    try:
        shutil.copy2(staged, temporary)
        os.replace(temporary, destination)
    finally:
        temporary.unlink(missing_ok=True)


def public_publish_preflight(staged: Path, destination: Path) -> str:
    if destination.exists():
        if sha256(destination) == sha256(staged):
            return "unchanged-identical"
        if not env_flag(ALLOW_PUBLIC_REPLACE_ENV):
            raise FileExistsError(
                f"Versioned public release already exists with different bytes: {destination}. "
                "Bump EDITION for a new release, or deliberately set "
                f"{ALLOW_PUBLIC_REPLACE_ENV}=1 to replace this edition."
            )
        mode = "replaced-by-explicit-override"
    else:
        mode = "created"
    return mode


def publish_public_immutable(staged: Path, destination: Path, *, expected_mode: str) -> None:
    current_mode = public_publish_preflight(staged, destination)
    if current_mode != expected_mode:
        raise RuntimeError(
            "Versioned public release changed after preflight: "
            f"expected {expected_mode}, found {current_mode}"
        )
    if current_mode != "unchanged-identical":
        publish(staged, destination)


def build() -> None:
    pandoc = find_tool("pandoc", "C0_PANDOC")
    soffice = find_tool("soffice", "C0_SOFFICE")
    pdfinfo = shutil.which("pdfinfo")
    pdffonts = shutil.which("pdffonts")

    lock_dir = PROJECT / "tmp" / "docs"
    lock_dir.mkdir(parents=True, exist_ok=True)
    lock_path = lock_dir / "student-textbook-build.lock"
    with lock_path.open("a+", encoding="utf-8") as lock_handle:
        try:
            fcntl.flock(lock_handle.fileno(), fcntl.LOCK_EX | fcntl.LOCK_NB)
        except BlockingIOError as exc:
            raise RuntimeError("Another student-textbook build is already running") from exc

        work_dir = Path(tempfile.mkdtemp(prefix="c0-student-textbook-build-"))
        try:
            markdown = assemble_source()
            source_qa = validate_source(markdown)
            staged_markdown = work_dir / "student-textbook.md"
            staged_markdown.write_text(markdown, encoding="utf-8")

            raw_docx = work_dir / "student-textbook-raw.docx"
            staged_docx = work_dir / DOCX_OUT.name
            pandoc_docx(pandoc, staged_markdown, raw_docx)
            style_qa = style_document(raw_docx, staged_docx)
            expected_worksheets = [
                str(record["title"]) for record in source_qa["worksheets"]
            ]
            if style_qa["worksheet_headings"] != expected_worksheets:
                raise ValueError(
                    "Styled worksheet registry differs from source registry: "
                    f"source={expected_worksheets}, styled={style_qa['worksheet_headings']}"
                )

            # A preliminary conversion supplies the page count for the DOCX
            # extended properties. Those properties do not affect layout; the
            # final conversion below verifies that the count remains stable.
            staged_pdf = convert_pdf(soffice, staged_docx, work_dir)
            preliminary_pages = len(PdfReader(str(staged_pdf)).pages)
            updated_docx_statistics = update_docx_extended_properties(
                staged_docx, pages=preliminary_pages
            )
            staged_pdf.unlink()
            staged_pdf = convert_pdf(soffice, staged_docx, work_dir)
            pdf_qa = validate_pdf(
                staged_pdf,
                pdfinfo,
                pdffonts,
                expected_worksheet_headings=expected_worksheets,
            )
            if pdf_qa["pages"] != preliminary_pages:
                raise ValueError(
                    "PDF page count changed after updating DOCX statistics: "
                    f"{preliminary_pages} -> {pdf_qa['pages']}"
                )
            docx_qa = validate_docx(
                staged_docx,
                expected_pages=int(pdf_qa["pages"]),
                expected_worksheet_headings=expected_worksheets,
            )
            if docx_qa["extended_statistics"] != {
                key: value
                for key, value in updated_docx_statistics.items()
                if key != "tables"
            }:
                raise ValueError("DOCX statistics changed unexpectedly during validation")

            palette_qa = validate_palette_contrast()
            pdf_digest = sha256(staged_pdf)
            pdf_bytes = staged_pdf.stat().st_size
            public_publish_mode = public_publish_preflight(staged_pdf, PUBLIC_PDF_OUT)

            manifest = {
                "schema": 2,
                "title": TITLE,
                "edition": EDITION,
                "built_at_utc": datetime.now(timezone.utc).isoformat(),
                "corpus_baseline": CORPUS_BASELINE,
                "palette": {
                    "burgundy": f"#{BURGUNDY}",
                    "sandstone": f"#{SANDSTONE}",
                    "espresso": f"#{ESPRESSO}",
                    "umber": f"#{UMBER}",
                    "pale_gold": f"#{PALE_GOLD}",
                },
                "sources": {
                    str(path.relative_to(PROJECT)): sha256(path)
                    for path in (FRONTMATTER, *CHAPTER_SOURCES)
                },
                "corpus_sources": {
                    str(path.relative_to(PROJECT)): sha256(path)
                    for path in CORPUS_SOURCES
                },
                "outputs": {
                    str(ASSEMBLED_SOURCE.relative_to(PROJECT)): sha256(staged_markdown),
                    str(DOCX_OUT.relative_to(PROJECT)): sha256(staged_docx),
                    str(PDF_OUT.relative_to(PROJECT)): pdf_digest,
                    str(PUBLIC_PDF_OUT.relative_to(PROJECT)): pdf_digest,
                },
                "release": {
                    "edition": EDITION,
                    "pdf": {
                        "public_path": str(PUBLIC_PDF_OUT.relative_to(PROJECT)),
                        "public_url": "/" + str(PUBLIC_PDF_OUT.relative_to(PROJECT / "public")),
                        "pages": pdf_qa["pages"],
                        "bytes": pdf_bytes,
                        "sha256": pdf_digest,
                    },
                    "public_publish_mode": public_publish_mode,
                    "replacement_override_env": ALLOW_PUBLIC_REPLACE_ENV,
                },
                "qa": {
                    "source": source_qa,
                    "styling": style_qa,
                    "docx": docx_qa,
                    "pdf": pdf_qa,
                    "palette_contrast": palette_qa,
                },
                "tools": {
                    "python": sys.version.splitlines()[0],
                    "pandoc": tool_version(pandoc, "--version"),
                    "libreoffice": tool_version(soffice, "--version"),
                    "pypdf": __import__("pypdf").__version__,
                    "python_docx": __import__("docx").__version__,
                },
            }
            staged_manifest = work_dir / MANIFEST_OUT.name
            staged_manifest.write_text(
                json.dumps(manifest, indent=2, sort_keys=True) + "\n", encoding="utf-8"
            )

            publish(staged_markdown, ASSEMBLED_SOURCE)
            publish(staged_docx, DOCX_OUT)
            publish(staged_pdf, PDF_OUT)
            publish(staged_manifest, MANIFEST_OUT)
            publish_public_immutable(
                staged_pdf,
                PUBLIC_PDF_OUT,
                expected_mode=public_publish_mode,
            )

            print(ASSEMBLED_SOURCE)
            print(DOCX_OUT)
            print(PDF_OUT)
            print(PUBLIC_PDF_OUT)
            print(MANIFEST_OUT)
        finally:
            shutil.rmtree(work_dir, ignore_errors=True)
            fcntl.flock(lock_handle.fileno(), fcntl.LOCK_UN)
    lock_path.unlink(missing_ok=True)


if __name__ == "__main__":
    try:
        build()
    except Exception as exc:
        print(f"Build failed: {exc}", file=sys.stderr)
        raise
